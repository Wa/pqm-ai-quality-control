"""Post-processing helpers for enterprise standard comparisons."""
from __future__ import annotations

import json
import os
import re
import shutil
import tempfile
import time
from datetime import datetime
from typing import List

import streamlit as st

from config import CONFIG
from ollama import Client as OllamaClient

from util import resolve_ollama_host

from tabs.shared import estimate_tokens, log_llm_metrics, report_exception, stream_text


def persist_compare_outputs(initial_dir: str, name_no_ext: str, prompt_texts: List[str], full_out_text: str) -> None:
    """Write prompt_{file}.txt and response_{file}.txt under initial_results."""

    try:
        total_parts = len(prompt_texts)
        prompt_out_lines = []
        for idx_p, ptxt in enumerate(prompt_texts, start=1):
            prompt_out_lines.append(f"提示词（第{idx_p}部分，共{total_parts}部分）：")
            prompt_out_lines.append(ptxt or "")
            prompt_out_lines.append("")
        prompt_text = "\n".join(prompt_out_lines).strip()
        prompt_path = os.path.join(initial_dir, f"prompt_{name_no_ext}.txt")
        with open(prompt_path, "w", encoding="utf-8") as handle:
            handle.write(prompt_text)

        response_path = os.path.join(initial_dir, f"response_{name_no_ext}.txt")
        with open(response_path, "w", encoding="utf-8") as handle:
            handle.write(full_out_text or "")
    except Exception as error:  # pragma: no cover - defensive UI feedback
        report_exception("保存比对过程失败", error)


def summarize_with_ollama(
    initial_dir: str,
    enterprise_out: str,
    session_id: str,
    name_no_ext: str,
    full_out_text: str,
) -> None:
    try:
        part_files: list[str] = []
        try:
            for fname in os.listdir(initial_dir):
                if fname.startswith(f"prompted_response_{name_no_ext}_pt") and fname.endswith(".txt"):
                    part_files.append(fname)
        except Exception as error:
            report_exception("读取Ollama提示列表失败", error, level="warning")
        part_files.sort(key=lambda value: value.lower())
        total_parts = len(part_files)
        host = resolve_ollama_host("ollama_9")
        ollama_client = OllamaClient(host=host)
        model = CONFIG["llm"].get("ollama_model")
        temperature = 0.7
        top_p = 0.9
        top_k = 40
        repeat_penalty = 1.1
        num_ctx = 40001
        num_thread = 4
        for part_idx, pfname in enumerate(part_files, start=1):
            p_path = os.path.join(initial_dir, pfname)
            try:
                with open(p_path, "r", encoding="utf-8") as handle:
                    prompt_text_all = handle.read()
            except Exception as error:
                report_exception(f"读取汇总提示失败({pfname})", error, level="warning")
                prompt_text_all = ""

            base = pfname[len("prompted_response_") :]
            json_name = f"json_{base}"
            json_path = os.path.join(initial_dir, json_name)
            existing_json_text = ""
            if os.path.isfile(json_path):
                try:
                    with open(json_path, "r", encoding="utf-8") as handle:
                        existing_json_text = handle.read()
                except Exception as error:
                    report_exception(f"读取已存在的JSON汇总失败({json_name})", error, level="warning")
                    existing_json_text = ""

            col_prompt2, col_resp2 = st.columns([1, 1])
            with col_prompt2:
                st.markdown(f"生成汇总表格提示词（第{part_idx}部分，共{total_parts}部分）")
                prompt_container2 = st.container(height=400)
                with prompt_container2:
                    with st.chat_message("user"):
                        placeholder = st.empty()
                        stream_text(placeholder, prompt_text_all, render_method="text")
                st.chat_input(
                    placeholder="",
                    disabled=True,
                    key=f"enterprise_prompted_prompt_{session_id}_{pfname}",
                )
            with col_resp2:
                st.markdown(f"生成汇总表格结果（第{part_idx}部分，共{total_parts}部分）")
                resp_container2 = st.container(height=400)
                with resp_container2:
                    with st.chat_message("assistant"):
                        ph_resp2 = st.empty()
                        if (existing_json_text or "").strip():
                            stream_text(ph_resp2, existing_json_text, render_method="write")
                            st.caption("已载入之前生成的JSON结果，无需重新调用模型。")
                        else:
                            response_text = ""
                            start_ts = time.time()
                            last_stats = None
                            error_msg = ""
                            chunks_seen = 0
                            try:
                                for chunk in ollama_client.chat(
                                    model=model,
                                    messages=[{"role": "user", "content": prompt_text_all}],
                                    stream=True,
                                    options={
                                        "temperature": temperature,
                                        "top_p": top_p,
                                        "top_k": top_k,
                                        "repeat_penalty": repeat_penalty,
                                        "num_ctx": num_ctx,
                                        "num_thread": num_thread,
                                    },
                                ):
                                    chunks_seen += 1
                                    new_text = chunk.get("message", {}).get("content", "")
                                    response_text += new_text
                                    ph_resp2.write(response_text)
                                    last_stats = chunk.get("eval_info") or chunk.get("stats") or last_stats
                            except Exception as error:  # pragma: no cover - runtime safeguard
                                error_msg = str(error)[:300]
                            finally:
                                dur_ms = int((time.time() - start_ts) * 1000)
                                prompt_tokens = (last_stats or {}).get("prompt_eval_count")
                                output_tokens = (last_stats or {}).get("eval_count")
                                if not error_msg and chunks_seen == 0:
                                    error_msg = "no_stream_chunks"
                                log_llm_metrics(
                                    enterprise_out,
                                    session_id,
                                    {
                                        "ts": datetime.now().isoformat(timespec="seconds"),
                                        "engine": "ollama",
                                        "model": model,
                                        "session_id": "",
                                        "file": name_no_ext,
                                        "part": part_idx,
                                        "phase": "summarize",
                                        "prompt_chars": len(prompt_text_all or ""),
                                        "prompt_tokens": prompt_tokens
                                        if isinstance(prompt_tokens, int)
                                        else estimate_tokens(prompt_text_all or ""),
                                        "output_chars": len(response_text or ""),
                                        "output_tokens": output_tokens
                                        if isinstance(output_tokens, int)
                                        else estimate_tokens(response_text or ""),
                                        "duration_ms": dur_ms,
                                        "success": 1 if (response_text or "").strip() else 0,
                                        "error": error_msg,
                                    },
                                )
                                try:
                                    with tempfile.NamedTemporaryFile(
                                        "w",
                                        delete=False,
                                        encoding="utf-8",
                                        dir=initial_dir,
                                    ) as tmp:
                                        tmp.write(response_text)
                                        tmp_name = tmp.name
                                    shutil.move(tmp_name, json_path)
                                except Exception as error:  # pragma: no cover - file IO failures
                                    report_exception("写入JSON汇总失败", error, level="warning")
    except Exception as error:  # pragma: no cover - global safeguard
        report_exception("调用Ollama生成汇总失败", error)


def aggregate_outputs(initial_dir: str, enterprise_out: str, session_id: str) -> None:
    final_dir = os.path.join(enterprise_out, "final_results")
    os.makedirs(final_dir, exist_ok=True)
    try:
        json_files = [
            os.path.join(initial_dir, fname)
            for fname in os.listdir(initial_dir)
            if fname.startswith("json_") and fname.lower().endswith(".txt")
        ]
    except Exception as error:
        report_exception("读取初始结果目录失败", error, level="warning")
        json_files = []

    columns = ["技术文件名", "技术文件内容", "企业标准", "不一致之处", "理由"]
    rows = []
    try:
        import pandas as pd  # type: ignore
    except ImportError as error:
        pd = None  # type: ignore[assignment]
        st.info(f"未安装 pandas，Excel 导出将被跳过：{error}")

    def _strip_code_fences(value: str) -> str:
        value = (value or "").strip()
        if value.startswith("```") and value.endswith("```"):
            value = value[3:-3].strip()
            if value.lower().startswith("json"):
                value = value[4:].strip()
        return value

    def _extract_json_text(text: str) -> str:
        s = (text or "").strip().lstrip("\ufeff")
        s = _strip_code_fences(s)
        try:
            json.loads(s)
            return s
        except json.JSONDecodeError:
            pass
        try:
            start = s.find("[")
            end = s.rfind("]")
            if start != -1 and end != -1 and end > start:
                candidate = s[start : end + 1]
                try:
                    json.loads(candidate)
                    return candidate
                except json.JSONDecodeError:
                    merged = re.sub(r"\]\s*,?\s*\[", ",", candidate)
                    json.loads(merged)
                    return merged
        except json.JSONDecodeError:
            pass
        try:
            start = s.find("{")
            end = s.rfind("}")
            if start != -1 and end != -1 and end > start:
                candidate = s[start : end + 1]
                json.loads(candidate)
                return candidate
        except json.JSONDecodeError:
            pass
        return s

    for jf in json_files:
        try:
            with open(jf, "r", encoding="utf-8") as handle:
                raw = handle.read()
        except Exception as error:
            report_exception(f"读取JSON失败({os.path.basename(jf)})", error, level="warning")
            continue
        parsed_text = _extract_json_text(raw)
        try:
            data = json.loads(parsed_text)
        except json.JSONDecodeError:
            continue
        if isinstance(data, dict):
            data = [data]
        if not isinstance(data, list):
            continue
        for row in data:
            if not isinstance(row, dict):
                continue
            rows.append(
                [
                    row.get("technical_file_name", ""),
                    row.get("technical_file_content", ""),
                    row.get("enterprise_standard", ""),
                    row.get("inconsistency", ""),
                    row.get("reason", ""),
                ]
            )

    if rows:
        csv_path = os.path.join(final_dir, f"企标检查对比结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
        try:
            with open(csv_path, "w", encoding="utf-8-sig", newline="") as handle:
                import csv

                writer = csv.writer(handle)
                writer.writerow(columns)
                writer.writerows(rows)
        except Exception as error:
            report_exception("写入CSV失败", error, level="warning")

        if "pd" in locals() and pd is not None:
            try:
                df = pd.DataFrame(rows, columns=columns)
                xlsx_path = os.path.join(final_dir, f"企标检查对比结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
                df.to_excel(xlsx_path, index=False)
            except Exception as error:
                report_exception("写入Excel失败", error, level="warning")

    try:
        pairs = []
        for fname in os.listdir(initial_dir):
            if fname.startswith("prompt_") and fname.endswith(".txt"):
                base = fname[len("prompt_") : -4]
                resp_name = f"response_{base}.txt"
                resp_path = os.path.join(initial_dir, resp_name)
                if os.path.isfile(resp_path):
                    pairs.append((os.path.join(initial_dir, fname), resp_path, base))
        if not pairs:
            return
        try:
            from docx import Document  # type: ignore
        except ImportError:
            st.info("未安装 python-docx，暂无法生成Word文档。请安装 python-docx 后重试。")
            return
        doc = Document()
        try:
            styles = doc.styles
            styles["Normal"].font.name = "宋体"
            styles["Heading 1"].font.name = "宋体"
            styles["Heading 2"].font.name = "宋体"
        except Exception:
            pass
        doc.add_heading("企标检查全部分析过程", level=0)
        doc.add_heading("目录", level=1)
        for _, _, base in pairs:
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{base}")
        for prompt_path, response_path, base in pairs:
            doc.add_heading(f"以下是《{base}》根据企业标准检查的分析过程：", level=1)
            try:
                with open(prompt_path, "r", encoding="utf-8") as handle:
                    prompt_text = handle.read()
            except Exception as error:
                report_exception(f"读取提示词文件失败({os.path.basename(prompt_path)})", error, level="warning")
                prompt_text = ""

            def _to_plain_table(text: str) -> str:
                try:
                    rows_local = []
                    for match in re.finditer(r"<tr[\s\S]*?>([\s\S]*?)</tr>", text, flags=re.IGNORECASE):
                        row_html = match.group(1)
                        cells = re.findall(
                            r"<(?:td|th)[^>]*>([\s\S]*?)</(?:td|th)>",
                            row_html,
                            flags=re.IGNORECASE,
                        )
                        cells_clean = [re.sub(r"<[^>]+>", "", cell).strip() for cell in cells]
                        if cells_clean:
                            rows_local.append("\t".join(cells_clean))
                    if rows_local:
                        return "\n".join(rows_local)
                    return re.sub(r"<[^>]+>", " ", text)
                except Exception:
                    return text

            for para in (prompt_text or "").splitlines():
                line = para.strip()
                if not line:
                    continue
                if re.match(r"^提示词（第\d+部分，共\d+部分）：", line):
                    doc.add_heading(line, level=2)
                else:
                    doc.add_paragraph(_to_plain_table(line))
            try:
                with open(response_path, "r", encoding="utf-8") as handle:
                    response_text = handle.read()
            except Exception as error:
                report_exception(f"读取比对结果失败({os.path.basename(response_path)})", error, level="warning")
                response_text = ""
            for para in (response_text or "").splitlines():
                line = para.strip()
                if not line:
                    continue
                doc.add_paragraph(_to_plain_table(line))
        ts_doc = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = os.path.join(final_dir, f"企标检查分析过程_{ts_doc}.docx")
        doc.save(doc_path)
        try:
            with open(doc_path, "rb") as handle:
                st.download_button(
                    label="下载分析过程Word",
                    data=handle.read(),
                    file_name=os.path.basename(doc_path),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_docx_{session_id}",
                )
        except Exception:
            pass
    except Exception as error:
        st.error(f"生成分析过程Word失败：{error}")


__all__ = ["aggregate_outputs", "persist_compare_outputs", "summarize_with_ollama"]
