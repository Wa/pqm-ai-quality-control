"""Post-processing helpers for the special symbols workflow."""
from __future__ import annotations

import json
import os
import re
import shutil
import tempfile
import time
from datetime import datetime
from typing import List, Optional, Tuple

try:  # pragma: no cover - backend execution may not ship Streamlit
    import streamlit as st  # type: ignore
except Exception:  # pragma: no cover
    st = None  # type: ignore

from config import CONFIG
from ollama import Client as OllamaClient

from util import resolve_ollama_host

from tabs.shared import estimate_tokens, log_llm_metrics, report_exception, stream_text


def _st_available() -> bool:
    """Return True when Streamlit runtime is importable and active."""

    return st is not None


def persist_compare_outputs(initial_dir: str, name_no_ext: str, prompt_texts: List[str], full_out_text: str) -> None:
    """Write prompt_{file}.txt and response_{file}.txt under initial_results."""

    try:
        total_parts = len(prompt_texts)
        prompt_out_lines = []
        for idx_p, ptxt in enumerate(prompt_texts, start=1):
            prompt_out_lines.append(f"提示词（第{idx_p}部分，共{total_parts}部分）：")
            prompt_out_lines.append(ptxt or "")
            prompt_out_lines.append("")
        prompt_text = "\n".join(prompt_out_lines).strip()
        prompt_path = os.path.join(initial_dir, f"prompt_{name_no_ext}.txt")
        with open(prompt_path, "w", encoding="utf-8") as handle:
            handle.write(prompt_text)

        response_path = os.path.join(initial_dir, f"response_{name_no_ext}.txt")
        with open(response_path, "w", encoding="utf-8") as handle:
            handle.write(full_out_text or "")
    except Exception as error:  # pragma: no cover - defensive UI feedback
        report_exception("保存比对过程失败", error)


def summarize_with_ollama(
    initial_dir: str,
    enterprise_out: str,
    session_id: str,
    name_no_ext: str,
    full_out_text: str,
) -> None:
    try:
        original_name = name_no_ext
        prompt_lines = [
            "你是一个严谨的结构化信息抽取助手。",
            "\n请将以下内容转换为 JSON 数组，每个对象包含如下几个键：",
            "- 条目（即特征/特性/过程）",
            "- 基准文件名和工作表名（如果有工作表名信息）",
            "- 基准文件中的特殊特性分类（★、☆、/）",
            "- 待检查文件名和工作表名（如果有工作表名信息）",
            "- 待检查文件中的特殊特性分类（★、☆、/）",
            "\n要求：",
            "1) 仅输出严格的 JSON（UTF-8，无注释、无多余文本）；",
            "2) 若内容包含多处对比，按条目拆分为多条 JSON 对象；",
            "3) 若某处信息缺失，请以空字符串 \"\" 占位，不要编造；",
            "4) 尽量保留可用于追溯定位的原文线索（如文件名、SHEET名、页码等）于相应字段中。",
            "\n下面是需要转换为 JSON 的内容：\n\n",
        ]
        instruction = "".join(prompt_lines)
        text = full_out_text or ""

        max_thinks_per_part = 20
        max_chars_per_part = 8000
        think_indices: list[int] = []
        needle = "<think>"
        start = 0
        while True:
            pos = text.find(needle, start)
            if pos == -1:
                break
            think_indices.append(pos)
            start = pos + len(needle)

        boundaries: list[int] = []
        prev = 0
        count_since_prev = 0
        for pos in think_indices:
            count_since_prev += 1
            if count_since_prev >= max_thinks_per_part or (pos - prev) >= max_chars_per_part:
                boundaries.append(pos)
                prev = pos
                count_since_prev = 0

        parts: list[str] = []
        prev = 0
        for boundary in boundaries:
            parts.append(text[prev:boundary])
            prev = boundary
        parts.append(text[prev:])

        # Remove legacy prompted_response files for this document to avoid stale parts lingering.
        try:
            for fname in os.listdir(initial_dir):
                if fname.startswith(f"prompted_response_{name_no_ext}_pt") and fname.endswith(".txt"):
                    os.remove(os.path.join(initial_dir, fname))
        except Exception:
            # Best effort cleanup; continue writing new parts even if deletion fails.
            pass

        for index, part in enumerate(parts, start=1):
            cleaned_part = re.sub(r"<think>[\s\S]*?</think>", "", part)
            prompted_path = os.path.join(
                initial_dir, f"prompted_response_{name_no_ext}_pt{index}.txt"
            )
            with open(prompted_path, "w", encoding="utf-8") as handle:
                handle.write(instruction)
                handle.write(cleaned_part)
    except Exception as error:
        report_exception("生成Ollama汇总提示失败", error, level="warning")

    try:
        part_files: list[str] = []
        try:
            for fname in os.listdir(initial_dir):
                if fname.startswith(f"prompted_response_{name_no_ext}_pt") and fname.endswith(".txt"):
                    part_files.append(fname)
        except Exception as error:
            report_exception("读取Ollama提示列表失败", error, level="warning")
        part_files.sort(key=lambda value: value.lower())
        total_parts = len(part_files)
        local_model = CONFIG["llm"].get("ollama_model") or "gpt-oss:latest"
        cloud_model = "gpt-oss:20b-cloud"
        local_client: Optional[OllamaClient]
        cloud_client: Optional[OllamaClient]
        try:
            host = resolve_ollama_host("ollama_9")
            local_client = OllamaClient(host=host)
        except Exception as error:
            local_client = None
            report_exception("初始化本地 gpt-oss 客户端失败", error, level="warning")
        cloud_host = CONFIG["llm"].get("ollama_cloud_host")
        cloud_api_key = CONFIG["llm"].get("ollama_cloud_api_key")
        if cloud_host and cloud_api_key:
            try:
                cloud_client = OllamaClient(
                    host=cloud_host,
                    headers={"Authorization": f"Bearer {cloud_api_key}"},
                )
            except Exception as error:
                cloud_client = None
                report_exception("初始化云端 gpt-oss 客户端失败", error, level="warning")
        else:
            cloud_client = None
        if local_client is None and cloud_client is None:
            report_exception(
                "未能初始化任何 gpt-oss 客户端，跳过汇总生成",
                RuntimeError("no available ollama client"),
                level="warning",
            )
            return
        temperature = 0.7
        top_p = 0.9
        top_k = 40
        repeat_penalty = 1.1
        num_ctx = 40001
        num_thread = 4
        show_ui = _st_available()
        attempt_sequence: List[Tuple[str, Optional[OllamaClient], Optional[str]]] = [
            ("本地 gpt-oss", local_client, local_model),
            ("云端 gpt-oss", cloud_client, cloud_model),
        ]
        for part_idx, pfname in enumerate(part_files, start=1):
            p_path = os.path.join(initial_dir, pfname)
            try:
                with open(p_path, "r", encoding="utf-8") as handle:
                    prompt_text_all = handle.read()
            except Exception as error:
                report_exception(f"读取汇总提示失败({pfname})", error, level="warning")
                prompt_text_all = ""

            base = pfname[len("prompted_response_") :]
            json_name = f"json_{base}"
            json_path = os.path.join(initial_dir, json_name)
            existing_json_text = ""
            if os.path.isfile(json_path):
                try:
                    with open(json_path, "r", encoding="utf-8") as handle:
                        existing_json_text = handle.read()
                except Exception as error:
                    report_exception(f"读取已存在的JSON汇总失败({json_name})", error, level="warning")
                    existing_json_text = ""

            placeholder = None
            if show_ui:
                col_prompt2, col_resp2 = st.columns([1, 1])
                with col_prompt2:
                    st.markdown(f"生成汇总表格提示词（第{part_idx}部分，共{total_parts}部分）")
                    prompt_container2 = st.container(height=400)
                    with prompt_container2:
                        with st.chat_message("user"):
                            placeholder = st.empty()
                            stream_text(placeholder, prompt_text_all, render_method="text")
                    st.chat_input(
                        placeholder="",
                        disabled=True,
                        key=f"enterprise_prompted_prompt_{session_id}_{pfname}",
                    )
                resp_column = col_resp2
            else:
                resp_column = None

            response_placeholder = None
            if show_ui and resp_column is not None:
                with resp_column:
                    st.markdown(f"生成汇总表格结果（第{part_idx}部分，共{total_parts}部分）")
                    resp_container2 = st.container(height=400)
                    with resp_container2:
                        with st.chat_message("assistant"):
                            response_placeholder = st.empty()

            response_text = (existing_json_text or "").strip()
            error_msg = ""
            last_stats = None
            dur_ms = 0
            prompt_tokens = None
            output_tokens = None
            generated_new = False

            if response_text:
                if show_ui and response_placeholder is not None:
                    stream_text(response_placeholder, response_text, render_method="write")
                    st.caption("已载入之前生成的JSON结果，无需重新调用模型。")
            else:
                start_ts = time.time()
                used_model: Optional[str] = None
                last_error: Optional[Exception] = None
                response_text = ""
                last_stats = None
                generated_new = False
                error_msg = ""
                chunks_seen = 0

                for attempt_label, attempt_client, attempt_model in attempt_sequence:
                    if attempt_client is None or not attempt_model:
                        continue

                    chunks_seen = 0
                    try:
                        if show_ui and response_placeholder is not None:
                            response_text = ""
                            for chunk in attempt_client.chat(
                                model=attempt_model,
                                messages=[{"role": "user", "content": prompt_text_all}],
                                stream=True,
                                options={
                                    "temperature": temperature,
                                    "top_p": top_p,
                                    "top_k": top_k,
                                    "repeat_penalty": repeat_penalty,
                                    "num_ctx": num_ctx,
                                    "num_thread": num_thread,
                                },
                            ):
                                new_text = chunk.get("message", {}).get("content", "")
                                if new_text:
                                    response_text += new_text
                                chunks_seen += 1
                                response_placeholder.write(response_text)
                                last_stats = chunk.get("eval_info") or chunk.get("stats") or last_stats
                            generated_new = True if response_text else False
                        else:
                            result = attempt_client.chat(
                                model=attempt_model,
                                messages=[{"role": "user", "content": prompt_text_all}],
                                stream=False,
                                options={
                                    "temperature": temperature,
                                    "top_p": top_p,
                                    "top_k": top_k,
                                    "repeat_penalty": repeat_penalty,
                                    "num_ctx": num_ctx,
                                    "num_thread": num_thread,
                                },
                            )
                            if isinstance(result, dict):
                                response_text = (
                                    result.get("message", {}).get("content")
                                    or result.get("response")
                                    or ""
                                )
                                last_stats = result.get("eval_info") or result.get("stats")
                                generated_new = True if response_text else False
                        used_model = attempt_model
                        error_msg = ""
                        last_error = None
                        break
                    except Exception as error:  # pragma: no cover - runtime safeguard
                        last_error = error
                        report_exception(f"调用 {attempt_label} 生成汇总失败", error, level="warning")
                        error_msg = str(error)[:300]
                        response_text = ""
                        last_stats = None
                        generated_new = False
                        if show_ui and response_placeholder is not None:
                            response_placeholder.write(f"{attempt_label} 调用失败：{error}")
                        continue

                dur_ms = int((time.time() - start_ts) * 1000)
                if used_model is None and last_error is not None and not error_msg:
                    error_msg = str(last_error)[:300]
                if not error_msg and chunks_seen == 0 and not response_text:
                    error_msg = "no_stream_chunks"
                prompt_tokens = (last_stats or {}).get("prompt_eval_count") if last_stats else None
                output_tokens = (last_stats or {}).get("eval_count") if last_stats else None
                model_for_logging = used_model or local_model
                log_llm_metrics(
                    enterprise_out,
                    session_id,
                    {
                        "ts": datetime.now().isoformat(timespec="seconds"),
                        "engine": "ollama",
                        "model": model_for_logging,
                        "session_id": "",
                        "file": name_no_ext,
                        "part": part_idx,
                        "phase": "summarize",
                        "prompt_chars": len(prompt_text_all or ""),
                        "prompt_tokens": prompt_tokens
                        if isinstance(prompt_tokens, int)
                        else estimate_tokens(prompt_text_all or "", model_for_logging),
                        "output_chars": len(response_text or ""),
                        "output_tokens": output_tokens
                        if isinstance(output_tokens, int)
                        else estimate_tokens(response_text or "", model_for_logging),
                        "duration_ms": dur_ms,
                        "success": 1 if (response_text or "").strip() else 0,
                        "error": error_msg,
                    },
                )
                if generated_new and response_text:
                    try:
                        with tempfile.NamedTemporaryFile(
                            "w",
                            delete=False,
                            encoding="utf-8",
                            dir=initial_dir,
                        ) as tmp:
                            tmp.write(response_text)
                            tmp_name = tmp.name
                        shutil.move(tmp_name, json_path)
                    except Exception as error:  # pragma: no cover - file IO failures
                        report_exception("写入JSON汇总失败", error, level="warning")
    except Exception as error:  # pragma: no cover - global safeguard
        report_exception("调用Ollama生成汇总失败", error)


def aggregate_outputs(initial_dir: str, enterprise_out: str, session_id: str) -> bool:
    final_dir = os.path.join(enterprise_out, "final_results")
    os.makedirs(final_dir, exist_ok=True)
    try:
        json_files = [
            os.path.join(initial_dir, fname)
            for fname in os.listdir(initial_dir)
            if fname.startswith("json_") and fname.lower().endswith(".txt")
        ]
    except Exception as error:
        report_exception("读取初始结果目录失败", error, level="warning")
        json_files = []

    columns: list[str] = []
    rows: list[dict[str, str]] = []
    try:
        import pandas as pd  # type: ignore
    except ImportError as error:
        pd = None  # type: ignore[assignment]
        if _st_available():
            st.info(f"未安装 pandas，Excel 导出将被跳过：{error}")
        else:
            report_exception("未安装 pandas，Excel 导出将被跳过", error, level="warning")

    def _strip_code_fences(value: str) -> str:
        value = (value or "").strip()
        if value.startswith("```") and value.endswith("```"):
            value = value[3:-3].strip()
            if value.lower().startswith("json"):
                value = value[4:].strip()
        return value

    def _extract_json_text(text: str) -> str:
        s = (text or "").strip().lstrip("\ufeff")
        s = _strip_code_fences(s)
        try:
            json.loads(s)
            return s
        except json.JSONDecodeError:
            pass
        try:
            start = s.find("[")
            end = s.rfind("]")
            if start != -1 and end != -1 and end > start:
                candidate = s[start : end + 1]
                try:
                    json.loads(candidate)
                    return candidate
                except json.JSONDecodeError:
                    merged = re.sub(r"\]\s*,?\s*\[", ",", candidate)
                    json.loads(merged)
                    return merged
        except json.JSONDecodeError:
            pass
        try:
            start = s.find("{")
            end = s.rfind("}")
            if start != -1 and end != -1 and end > start:
                candidate = s[start : end + 1]
                json.loads(candidate)
                return candidate
        except json.JSONDecodeError:
            pass
        return s

    for jf in json_files:
        try:
            with open(jf, "r", encoding="utf-8") as handle:
                raw = handle.read()
        except Exception as error:
            report_exception(f"读取JSON失败({os.path.basename(jf)})", error, level="warning")
            continue

        parsed_text = _extract_json_text(raw)
        try:
            data = json.loads(parsed_text)
        except json.JSONDecodeError:
            continue
        if isinstance(data, dict):
            data = [data]
        if not isinstance(data, list):
            continue

        def _stringify(value: object) -> str:
            if value is None:
                return ""
            if isinstance(value, (dict, list)):
                try:
                    return json.dumps(value, ensure_ascii=False)
                except Exception:
                    return str(value)
            return str(value)

        for row in data:
            normalized: dict[str, str] = {}
            if isinstance(row, dict):
                items = row.items()
            else:
                items = [("值", row)]
            for key, value in items:
                column_name = str(key).strip() or "值"
                value_str = _stringify(value)
                normalized[column_name] = value_str
                if column_name not in columns:
                    columns.append(column_name)
            rows.append(normalized)

    # Filter out rows with matching symbols
    if rows:
        # Lazy import to avoid circular dependency
        from .background import _filter_identical_matches

        filtered_rows, removed_count = _filter_identical_matches(rows)
        rows = filtered_rows
        # Note: removed_count rows with matching symbols have been filtered out

    csv_path: str | None = None
    xlsx_path: str | None = None

    if rows and columns:
        table_rows = [{column: row.get(column, "") for column in columns} for row in rows]
        csv_path = os.path.join(
            final_dir,
            f"特殊特性符号检查对比结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
        )
        try:
            with open(csv_path, "w", encoding="utf-8-sig", newline="") as handle:
                import csv

                writer = csv.DictWriter(handle, fieldnames=columns)
                writer.writeheader()
                writer.writerows(table_rows)
        except Exception as error:
            report_exception("写入CSV失败", error, level="warning")

        if "pd" in locals() and pd is not None:
            try:
                df = pd.DataFrame(table_rows, columns=columns)
                xlsx_path = os.path.join(
                    final_dir,
                    f"特殊特性符号检查对比结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                )
                df.to_excel(xlsx_path, index=False)
            except Exception as error:
                report_exception("写入Excel失败", error, level="warning")

        if _st_available():
            try:
                if csv_path and os.path.exists(csv_path):
                    with open(csv_path, "rb") as handle:
                        st.download_button(
                            label="下载CSV结果",
                            data=handle.read(),
                            file_name=os.path.basename(csv_path),
                            mime="text/csv",
                            key=f"download_csv_{session_id}",
                        )
                if xlsx_path and os.path.exists(xlsx_path):
                    with open(xlsx_path, "rb") as handle:
                        st.download_button(
                            label="下载Excel结果",
                            data=handle.read(),
                            file_name=os.path.basename(xlsx_path),
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"download_xlsx_{session_id}",
                        )
            except Exception as error:
                report_exception("生成下载链接失败", error, level="warning")

    has_rows = bool(rows)

    try:
        pairs = []
        for fname in os.listdir(initial_dir):
            if fname.startswith("prompt_") and fname.endswith(".txt"):
                base = fname[len("prompt_") : -4]
                resp_name = f"response_{base}.txt"
                resp_path = os.path.join(initial_dir, resp_name)
                if os.path.isfile(resp_path):
                    pairs.append((os.path.join(initial_dir, fname), resp_path, base))
        if not pairs:
            return has_rows
        try:
            from docx import Document  # type: ignore
        except ImportError as error:
            if _st_available():
                st.info("未安装 python-docx，暂无法生成Word文档。请安装 python-docx 后重试。")
            else:
                report_exception("缺少 python-docx 依赖，跳过生成Word文档", error, level="warning")
            return has_rows
        doc = Document()
        try:
            styles = doc.styles
            styles["Normal"].font.name = "宋体"
            styles["Heading 1"].font.name = "宋体"
            styles["Heading 2"].font.name = "宋体"
        except Exception:
            pass
        doc.add_heading("特殊特性符号检查全部分析过程", level=0)
        doc.add_heading("目录", level=1)
        for _, _, base in pairs:
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{base}")
        for prompt_path, response_path, base in pairs:
            doc.add_heading(f"以下是《{base}》根据基准文件检查的分析过程：", level=1)
            try:
                with open(prompt_path, "r", encoding="utf-8") as handle:
                    prompt_text = handle.read()
            except Exception as error:
                report_exception(f"读取提示词文件失败({os.path.basename(prompt_path)})", error, level="warning")
                prompt_text = ""

            def _to_plain_table(text: str) -> str:
                try:
                    rows_local = []
                    for match in re.finditer(r"<tr[\s\S]*?>([\s\S]*?)</tr>", text, flags=re.IGNORECASE):
                        row_html = match.group(1)
                        cells = re.findall(
                            r"<(?:td|th)[^>]*>([\s\S]*?)</(?:td|th)>",
                            row_html,
                            flags=re.IGNORECASE,
                        )
                        cells_clean = [re.sub(r"<[^>]+>", "", cell).strip() for cell in cells]
                        if cells_clean:
                            rows_local.append("\t".join(cells_clean))
                    if rows_local:
                        return "\n".join(rows_local)
                    return re.sub(r"<[^>]+>", " ", text)
                except Exception:
                    return text

            for para in (prompt_text or "").splitlines():
                line = para.strip()
                if not line:
                    continue
                if re.match(r"^提示词（第\d+部分，共\d+部分）：", line):
                    doc.add_heading(line, level=2)
                else:
                    doc.add_paragraph(_to_plain_table(line))
            try:
                with open(response_path, "r", encoding="utf-8") as handle:
                    response_text = handle.read()
            except Exception as error:
                report_exception(f"读取比对结果失败({os.path.basename(response_path)})", error, level="warning")
                response_text = ""
            for para in (response_text or "").splitlines():
                line = para.strip()
                if not line:
                    continue
                doc.add_paragraph(_to_plain_table(line))
        ts_doc = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_path = os.path.join(final_dir, f"特殊特性符号检查分析过程_{ts_doc}.docx")
        doc.save(doc_path)
        if _st_available():
            try:
                with open(doc_path, "rb") as handle:
                    st.download_button(
                        label="下载分析过程Word",
                        data=handle.read(),
                        file_name=os.path.basename(doc_path),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_docx_{session_id}",
                    )
            except Exception:
                pass
    except Exception as error:
        report_exception("生成分析过程Word失败", error, level="warning")

    return has_rows

__all__ = ["aggregate_outputs", "persist_compare_outputs", "summarize_with_ollama"]
