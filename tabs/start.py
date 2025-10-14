if st.session_state.get(f"enterprise_running_{session_id}"):
	# Retrieve context saved before rerun
	std_txt_files = st.session_state.get(f"enterprise_std_txt_files_{session_id}") or []
	exam_txt_files = st.session_state.get(f"enterprise_exam_txt_files_{session_id}") or []
	enterprise_out = st.session_state.get(f"enterprise_out_root_{session_id}") or enterprise_out_root
	std_txt_dir = st.session_state.get(f"enterprise_standards_txt_dir_{session_id}") or standards_txt_dir
	exam_txt_dir = st.session_state.get(f"enterprise_examined_txt_dir_{session_id}") or examined_txt_dir

	# Upload standards once (optional)
	if std_txt_files:
		with st.status("Sync standards to KB...", expanded=False) as status:
			try:
				kb_name_dyn = f"{session_id}_{TAB_SLUG}"
				kid = find_knowledge_id_by_name(BISHENG_BASE_URL, BISHENG_API_KEY or None, kb_name_dyn)
				if not kid:
					kid = create_knowledge(BISHENG_BASE_URL, BISHENG_API_KEY or None, kb_name_dyn, model=str(KB_MODEL_ID))
				if kid:
					res = kb_sync_folder(
						base_url=BISHENG_BASE_URL,
						api_key=BISHENG_API_KEY or None,
						knowledge_id=int(kid),
						folder_path=std_txt_dir,
						clear_first=False,
						chunk_size=1000,
						chunk_overlap=0,
						separators=["\n\n", "\n"],
						separator_rule=["after", "after"],
					)
					status.update(label=f"KB sync: uploaded {len(res.get('uploaded', []))}, deleted {len(res.get('deleted', []))}, skipped {len(res.get('skipped', []))}", state="complete")
				else:
					status.update(label="KB create/lookup failed (check server auth)", state="error")
			except Exception as e:
				status.update(label=f"KB sync failed: {e}", state="error")

	# Iterate examined texts
	exam_txt_files.sort(key=lambda x: x.lower())
	bisheng_session_id = st.session_state.get(f"bisheng_session_{session_id}")
	initial_dir = os.path.join(enterprise_out, 'initial_results')
	os.makedirs(initial_dir, exist_ok=True)
	# 预热步骤：在并发开始前，串行对 Bisheng Flow 发起一次极短请求，促使检索/LLM 初始化与缓存
	# 说明：首次请求常见的冷启动（模型加载、连接池、检索索引唤醒）会导致首批并发请求失败率上升；
	# 通过一次轻量的预热，可以显著降低“第一批全挂”的概率。返回内容无需使用。
	try:
		if not st.session_state.get(f"enterprise_warmup_done_{session_id}"):
			warmup_prompt = "预热：请简短回复 'gotcha' 即可。"
			_ = call_flow_process(
				base_url=BISHENG_BASE_URL,
				flow_id=BISHENG_FLOW_ID,
				question=warmup_prompt,
				kb_id=None,
				input_node_id=FLOW_INPUT_NODE_ID,
				api_key=BISHENG_API_KEY or None,
				session_id=None,
				history_count=0,
				extra_tweaks={"CombineDocsChain-520ca": {"token_max": 5000}},
				milvus_node_id=FLOW_MILVUS_NODE_ID,
				es_node_id=FLOW_ES_NODE_ID,
				timeout_s=60,
				max_retries=0,
				# clear_cache=True,
			)
			st.session_state[f"enterprise_warmup_done_{session_id}"] = True
	except Exception:
		pass
	# Start fresh: original per-file splitting and live prompting
	for idx_file, name in enumerate(exam_txt_files, start=1):
		src_path = os.path.join(exam_txt_dir, name)
		st.markdown(f"**📄 正在比对第{idx_file}个文件，共{len(exam_txt_files)}个：{name}**")
		try:
			with open(src_path, 'r', encoding='utf-8') as f:
				doc_text = f.read()
		except Exception as e:
			st.error(f"读取失败：{e}")
			continue
		if not doc_text.strip():
			st.info("文件为空，跳过。")
			continue
		chunks = split_to_chunks(doc_text, int(BISHENG_MAX_WORDS))
		prompt_prefix = (
			"请作为企业标准符合性检查专家，审阅待检查文件与企业标准是否一致。"
			"以列表形式列出不一致的点，并引用原文证据（简短摘录）、标明出处（提供企业标准文件的文件名）。\n"
			"输出的内容要言简意赅，列出不一致的点即可，不需要列出一致的点，也不需要列出企业标准中缺失的点，最后不需要总结。\n"
			"由于待检查文件较长，我将分成多个部分将其上传给你。以下是待检查文件的一部分。\n"
		)
		full_out_text = ""
		prompt_texts = []
		for i, piece in enumerate(chunks, start=1):
			col_prompt, col_response = st.columns([1, 1])
			prompt_text = f"{prompt_prefix}{piece}"
			prompt_texts.append(prompt_text)
			with col_prompt:
				st.markdown(f"提示词（第{i}部分，共{len(chunks)}部分）")
				prompt_container = st.container(height=400)
				with prompt_container:
					with st.chat_message("user"):
						prompt_placeholder = st.empty()
						words = prompt_text.split()
						streamed = ""
						for j in range(0, len(words), 30):
							chunk_words = words[j:j+30]
							streamed += " ".join(chunk_words) + " "
							prompt_placeholder.text(streamed.strip())
					st.chat_input(placeholder="", disabled=True, key=f"enterprise_prompt_{session_id}_{idx_file}_{i}")
			with col_response:
				st.markdown(f"AI比对结果（第{i}部分，共{len(chunks)}部分）")
				response_container = st.container(height=400)
				with response_container:
					with st.chat_message("assistant"):
						response_placeholder = st.empty()
						try:
							# Determine per-user KB name and call Flow with tweaks for per-run KB binding
							kb_name_dyn = f"{session_id}_{TAB_SLUG}"
							kid = find_knowledge_id_by_name(BISHENG_BASE_URL, BISHENG_API_KEY or None, kb_name_dyn)
							start_ts = time.time()
							res = call_flow_process(
								base_url=BISHENG_BASE_URL,
								flow_id=BISHENG_FLOW_ID,
								question=prompt_text,
								kb_id=kid,
								input_node_id=FLOW_INPUT_NODE_ID,
								api_key=BISHENG_API_KEY or None,
								session_id=bisheng_session_id,
								history_count=0,
								extra_tweaks=None,
								milvus_node_id=FLOW_MILVUS_NODE_ID,
								es_node_id=FLOW_ES_NODE_ID,
								timeout_s=180,        # 新增
								max_retries=2,        # 新增
								# clear_cache=True,
								)
							ans_text, new_sid = parse_flow_answer(res)
							dur_ms = int((time.time() - start_ts) * 1000)
							try:
								from datetime import datetime as _dt
								_log_llm_metrics(
									enterprise_out,
									session_id,
									{
										"ts": _dt.now().isoformat(timespec="seconds"),
										"engine": "bisheng",
										"model": "qwen3",
										"session_id": bisheng_session_id or "",
										"file": name,
										"part": i,
										"phase": "compare",
										"prompt_chars": len(prompt_text or ""),
										"prompt_tokens": _estimate_tokens(prompt_text or ""),
										"output_chars": len(ans_text or ""),
										"output_tokens": _estimate_tokens(ans_text or ""),
										"duration_ms": dur_ms,
										"success": 1 if (ans_text or "").strip() else 0,
										"error": res.get("error") if isinstance(res, dict) else "",
									},
								)
							except Exception:
								pass
							if new_sid:
								bisheng_session_id = new_sid
								st.session_state[f"bisheng_session_{session_id}"] = bisheng_session_id
							response_placeholder.write(ans_text or "")
							full_out_text += ("\n\n" if full_out_text else "") + (ans_text or "")
						except Exception as e:
							response_placeholder.error(f"调用失败：{e}")
					st.chat_input(placeholder="", disabled=True, key=f"enterprise_response_{session_id}_{idx_file}_{i}")
		# Persist per-file combined output
		try:
			name_no_ext = os.path.splitext(name)[0]
			# Also persist all prompt parts as a single file: prompt_{name_no_ext}.txt
			try:
				total_parts = len(prompt_texts)
				prompt_out_lines = []
				for idx_p, ptxt in enumerate(prompt_texts, start=1):
					prompt_out_lines.append(f"提示词（第{idx_p}部分，共{total_parts}部分）：")
					prompt_out_lines.append(ptxt)
				prompt_out_text = "\n".join(prompt_out_lines)
				prompt_out_path = os.path.join(initial_dir, f"prompt_{name_no_ext}.txt")
				with open(prompt_out_path, 'w', encoding='utf-8') as pf:
					pf.write(prompt_out_text)
			except Exception:
				pass
			out_path = os.path.join(initial_dir, f"response_{name_no_ext}.txt")
			with open(out_path, 'w', encoding='utf-8') as outf:
				outf.write(full_out_text)
			# Immediately produce chunked prompted_response_*_ptN.txt files
			try:
				original_name = name_no_ext
				prompt_lines = [
					"你是一个严谨的结构化信息抽取助手。以下内容来自一个基于 RAG 的大语言模型 (RAG-LLM) 系统，",
					"用于将若干技术文件与企业标准进行逐条比对，输出发现的不符合项及其理由。本文件对应的原始技术文件名称为：",
					f"{original_name}。",
					"\n请将随后的全文内容转换为 JSON 数组（list of objects），每个对象包含如下五个键：",
					f"- 技术文件名：\"{original_name}\"",
					"- 技术文件内容：与该条不一致点相关的技术文件条目或段落名称/编号/摘要",
					"- 企业标准：被对比的企业标准条款名称/编号/摘要",
					"- 不一致之处：不符合或存在差异的具体点，尽量具体明确",
					"- 理由：判断不一致的依据与简要解释（保持客观、可追溯）",
					"\n要求：",
					"1) 仅输出严格的 JSON（UTF-8，无注释、无多余文本），键名使用上述中文；",
					"2) 若内容包含多处对比，按条目拆分为多条 JSON 对象；",
					"3) 若某处信息缺失，请以空字符串 \"\" 占位，不要编造；",
					"4) 尽量保留可用于追溯定位的原文线索（如编号、标题、页码等）于相应字段中。",
					"\n下面是需要转换为 JSON 的原始比对输出：\n\n",
				]
				instr = "".join(prompt_lines)
				# Split content by every 20 occurrences of "<think>" markers
				text = full_out_text or ""
				# 更稳健的切分：同时限制每段最多包含的 <think> 数量与最大字符数，避免单段过大
				max_thinks_per_part = 20
				max_chars_per_part = 8000
				think_indices = []
				needle = "<think>"
				start = 0
				while True:
					pos = text.find(needle, start)
					if pos == -1:
						break
					think_indices.append(pos)
					start = pos + len(needle)
				boundaries = []
				prev = 0
				count_since_prev = 0
				for pos in think_indices:
					count_since_prev += 1
					if count_since_prev >= max_thinks_per_part or (pos - prev) >= max_chars_per_part:
						boundaries.append(pos)
						prev = pos
						count_since_prev = 0
				parts = []
				prev = 0
				for b in boundaries:
					parts.append(text[prev:b])
					prev = b
				parts.append(text[prev:])
				# Write each part to prompted_response_<name>_ptN.txt after removing <think> blocks
				for idx_part, part in enumerate(parts, start=1):
					cleaned_part = re.sub(r"<think>[\s\S]*?</think>", "", part)
					prompted_path = os.path.join(initial_dir, f"prompted_response_{name_no_ext}_pt{idx_part}.txt")
					with open(prompted_path, 'w', encoding='utf-8') as pf:
						pf.write(instr)
						pf.write(cleaned_part)
			except Exception:
				pass

			# For each prompted_response_* part, call Ollama LLM and display side-by-side
			try:
				# Discover all parts for current file
				part_files = []
				try:
					for fn in os.listdir(initial_dir):
						if fn.startswith(f"prompted_response_{name_no_ext}_pt") and fn.endswith('.txt'):
							part_files.append(fn)
				except Exception:
					part_files = []
				part_files.sort(key=lambda x: x.lower())
				total_parts = len(part_files)
				# Prepare LLM client (Ollama on 10.31.60.9 by default)
				host = resolve_ollama_host("ollama_9")
				ollama_client = OllamaClient(host=host)
				model = CONFIG["llm"].get("ollama_model")
				temperature = 0.7
				top_p = 0.9
				top_k = 40
				repeat_penalty = 1.1
				num_ctx = 40001
				num_thread = 4
				for part_idx, pfname in enumerate(part_files, start=1):
					p_path = os.path.join(initial_dir, pfname)
					try:
						with open(p_path, 'r', encoding='utf-8') as fpr:
							prompt_text_all = fpr.read()
					except Exception:
						prompt_text_all = ""
					# Two-column display for this prompt-response pair
					col_prompt2, col_resp2 = st.columns([1, 1])
					with col_prompt2:
						st.markdown(f"生成汇总表格提示词（第{part_idx}部分，共{total_parts}部分）")
						prompt_container2 = st.container(height=400)
						with prompt_container2:
							with st.chat_message("user"):
								ph2 = st.empty()
								# Stream display the prompt text in chunks of words
								words = (prompt_text_all or "").split()
								streamed2 = ""
								for idxw in range(0, len(words), 30):
									chunk_words = words[idxw:idxw+30]
									streamed2 += " ".join(chunk_words) + " "
									ph2.text(streamed2.strip())
							st.chat_input(placeholder="", disabled=True, key=f"enterprise_prompted_prompt_{session_id}_{pfname}")
					with col_resp2:
						st.markdown(f"生成汇总表格结果（第{part_idx}部分，共{total_parts}部分）")
						resp_container2 = st.container(height=400)
						with resp_container2:
							with st.chat_message("assistant"):
								ph_resp2 = st.empty()
								response_text = ""
								start_ts = time.time()
								last_stats = None
								error_msg = ""
								chunks_seen = 0
								try:
									for chunk in ollama_client.chat(
										model=model,
										messages=[{"role": "user", "content": prompt_text_all}],
										stream=True,
										options={ "temperature": temperature, "top_p": top_p, "top_k": top_k,
										          "repeat_penalty": repeat_penalty, "num_ctx": num_ctx, "num_thread": num_thread }
									):
										chunks_seen += 1
										new_text = chunk.get('message', {}).get('content', '')
										response_text += new_text
										ph_resp2.write(response_text)
										# 某些引擎会在最后一个分块里返回统计信息
										last_stats = chunk.get('eval_info') or chunk.get('stats') or last_stats
								except Exception as e:
									error_msg = str(e)[:300]
								finally:
									from datetime import datetime as _dt
									dur_ms = int((time.time() - start_ts) * 1000)
									prompt_tokens = (last_stats or {}).get('prompt_eval_count')
									output_tokens = (last_stats or {}).get('eval_count')
									if not error_msg and chunks_seen == 0:
										error_msg = "no_stream_chunks"
									_log_llm_metrics(
										enterprise_out, session_id,
										{
											"ts": _dt.now().isoformat(timespec="seconds"),
											"engine": "ollama",
											"model": model,
											"session_id": "",
											"file": name_no_ext,
											"part": part_idx,
											"phase": "summarize",
											"prompt_chars": len(prompt_text_all or ""),
											"prompt_tokens": prompt_tokens if isinstance(prompt_tokens, int) else _estimate_tokens(prompt_text_all or ""),
											"output_chars": len(response_text or ""),
											"output_tokens": output_tokens if isinstance(output_tokens, int) else _estimate_tokens(response_text or ""),
											"duration_ms": dur_ms,
											"success": 1 if (response_text or "").strip() else 0,
											"error": "error_msg"
										}
									)

								# Save LLM response to json_<original>_ptN.txt
								try:
									# Map prompted_response_X_ptN.txt -> json_X_ptN.txt
									base = pfname[len("prompted_response_"):]
									json_name = f"json_{base}"
									json_path = os.path.join(initial_dir, json_name)
									with open(json_path, 'w', encoding='utf-8') as jf:
										jf.write(response_text)
								except Exception:
									pass
			except Exception:
				pass
		except Exception as e:
			st.error(f"保存结果失败：{e}")

	# End of current run; clear running flag (no final aggregation here)
	try:
		st.session_state[f"enterprise_running_{session_id}"] = False
	except Exception as e:
		st.error(f"流程收尾失败：{e}")

	# After LLM step: aggregate json_*_ptN.txt into CSV and XLSX in final_results
	try:
		final_dir = os.path.join(enterprise_out, 'final_results')
		os.makedirs(final_dir, exist_ok=True)
		# Collect json_*.txt under initial_results
		json_files = []
		try:
			for fn in os.listdir(initial_dir):
				if fn.startswith('json_') and fn.lower().endswith('.txt'):
					json_files.append(os.path.join(initial_dir, fn))
		except Exception:
			json_files = []
		# Build rows
		columns = ["技术文件名", "技术文件内容", "企业标准", "不一致之处", "理由"]
		rows = []
		import csv
		import pandas as pd
		def _strip_code_fences(s: str) -> str:
			s = (s or "").strip()
			if s.startswith("```") and s.endswith("```"):
				s = s[3:-3].strip()
				if s.lower().startswith("json"):
					s = s[4:].strip()
			return s
		def _extract_json_text(text: str) -> str:
			s = (text or "").strip().lstrip("\ufeff")
			s = _strip_code_fences(s)
			# try direct
			try:
				json.loads(s)
				return s
			except Exception:
				pass
			# try array
			try:
				start = s.find("[")
				end = s.rfind("]")
				if start != -1 and end != -1 and end > start:
					candidate = s[start:end+1]
					try:
						json.loads(candidate)
						return candidate
					except Exception:
						merged = re.sub(r"\]\s*,?\s*\[", ",", candidate)
						json.loads(merged)
						return merged
			except Exception:
				pass
			# try object
			try:
				start = s.find("{")
				end = s.rfind("}")
				if start != -1 and end != -1 and end > start:
					candidate = s[start:end+1]
					json.loads(candidate)
					return candidate
			except Exception:
				pass
			return s
		for jf in json_files:
			try:
				with open(jf, 'r', encoding='utf-8') as f:
					text = f.read()
			except Exception:
				text = ""
			js = _extract_json_text(text)
			try:
				parsed = json.loads(js)
			except Exception:
				parsed = None
			items = []
			if isinstance(parsed, list):
				items = [x for x in parsed if isinstance(x, dict)]
			elif isinstance(parsed, dict):
				vals = list(parsed.values())
				for v in vals:
					if isinstance(v, list) and all(isinstance(i, dict) for i in v):
						items = v
						break
				if not items:
					items = [parsed]
			# Derive original source filename from json_*.txt → strip 'json_' and trailing _ptN.txt
			base_name = os.path.basename(jf)
			orig_name = base_name[5:] if base_name.startswith('json_') else base_name
			try:
				orig_name = re.sub(r"_pt\d+\.txt$", "", orig_name)
			except Exception:
				pass
			for obj in items:
				row = [
					orig_name,
					str(obj.get("技术文件内容", "")),
					str(obj.get("企业标准", "")),
					str(obj.get("不一致之处", "")),
					str(obj.get("理由", "")),
				]
				rows.append(row)
		# Write CSV
		from datetime import datetime as _dt
		ts = _dt.now().strftime('%Y%m%d_%H%M%S')
		csv_path = os.path.join(final_dir, f"企标检查结果_{ts}.csv")
		with open(csv_path, 'w', encoding='utf-8-sig', newline='') as cf:
			writer = csv.writer(cf)
			writer.writerow(columns)
			for r in rows:
				writer.writerow(r)
		# Write XLSX
		xlsx_path = os.path.join(final_dir, f"企标检查结果_{ts}.xlsx")
		try:
			df = pd.DataFrame(rows, columns=columns)
			df.to_excel(xlsx_path, index=False, engine='openpyxl')
		except Exception:
			pass
		# Offer downloads
		try:
			with open(csv_path, 'rb') as fcsv:
				st.download_button(label="下载CSV结果", data=fcsv.read(), file_name=os.path.basename(csv_path), mime='text/csv', key=f"download_csv_{session_id}")
			with open(xlsx_path, 'rb') as fxlsx:
				st.download_button(label="下载Excel结果", data=fxlsx.read(), file_name=os.path.basename(xlsx_path), mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', key=f"download_xlsx_{session_id}")
		except Exception:
			pass
		# Build Word document of analysis process from prompt_*.txt and response_*.txt
		try:
			try:
				from docx import Document
			except Exception:
				Document = None
			if Document is not None:
				# Pair prompt_/response_ files by base name
				pairs = []
				prompt_files = [f for f in os.listdir(initial_dir) if f.startswith('prompt_') and f.lower().endswith('.txt')]
				for pf in prompt_files:
					base = pf[len('prompt_'):-4]
					rf = os.path.join(initial_dir, f"response_{base}.txt")
					if os.path.isfile(rf):
						pairs.append((os.path.join(initial_dir, pf), rf, base))
				pairs.sort(key=lambda x: x[2].lower())
				if pairs:
					doc = Document()
					# Set default fonts to 宋体 where possible
					try:
						styles = doc.styles
						styles['Normal'].font.name = '宋体'
						# Heading styles
						styles['Heading 1'].font.name = '宋体'
						styles['Heading 2'].font.name = '宋体'
					except Exception:
						pass
					# Main title
					doc.add_heading("企标检查全部分析过程", level=0)
					# Simple TOC section (static list)
					doc.add_heading("目录", level=1)
					for _, _, base in pairs:
						p = doc.add_paragraph()
						p.add_run(f"{base}")
					for p_path, r_path, base in pairs:
						doc.add_heading(f"以下是《{base}》根据企业标准检查的分析过程：", level=1)
						# Insert prompt content (mark each 提示词... 行为标题2)
						try:
							with open(p_path, 'r', encoding='utf-8') as f:
								ptext = f.read()
						except Exception:
							ptext = ""
						def _to_plain_table(txt: str) -> str:
							# Very simple HTML table conversion: extract <tr> rows and <td>/<th> cells
							try:
								import re as _re
								rows = []
								for m in _re.finditer(r"<tr[\s\S]*?>([\s\S]*?)</tr>", txt, flags=_re.IGNORECASE):
									row_html = m.group(1)
									cells = _re.findall(r"<(?:td|th)[^>]*>([\s\S]*?)</(?:td|th)>", row_html, flags=_re.IGNORECASE)
									cells_clean = [ _re.sub(r"<[^>]+>", "", c).strip() for c in cells ]
									if cells_clean:
										rows.append("\t".join(cells_clean))
								if rows:
									return "\n".join(rows)
								# Fallback: strip tags
								return _re.sub(r"<[^>]+>", " ", txt)
							except Exception:
								return txt
						for para in (ptext or "").splitlines():
							line = para.strip()
							if not line:
								continue
							if re.match(r"^提示词（第\d+部分，共\d+部分）：", line):
								doc.add_heading(line, level=2)
							else:
								doc.add_paragraph(_to_plain_table(line))
						# Insert response content
						try:
							with open(r_path, 'r', encoding='utf-8') as f:
								rtext = f.read()
						except Exception:
							rtext = ""
						for para in (rtext or "").splitlines():
							line = para.strip()
							if not line:
								continue
							doc.add_paragraph(_to_plain_table(line))
					from datetime import datetime as _dt
					ts_doc = _dt.now().strftime('%Y%m%d_%H%M%S')
					doc_path = os.path.join(final_dir, f"企标检查分析过程_{ts_doc}.docx")
					doc.save(doc_path)
					# Optional download button
					try:
						with open(doc_path, 'rb') as fdoc:
							st.download_button(label="下载分析过程Word", data=fdoc.read(), file_name=os.path.basename(doc_path), mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', key=f"download_docx_{session_id}")
					except Exception:
						pass
			else:
				st.info("未安装 python-docx，暂无法生成Word文档。请安装 python-docx 后重试。")
		except Exception as e:
			st.error(f"生成分析过程Word失败：{e}")
	except Exception as e:
		st.error(f"汇总导出失败：{e}")
# The end