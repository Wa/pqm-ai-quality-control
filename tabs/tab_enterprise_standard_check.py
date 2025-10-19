import streamlit as st
import os
import io
import zipfile
import json
import re
import time
import requests
from util import ensure_session_dirs, handle_file_upload, resolve_ollama_host
from config import CONFIG
from ollama import Client as OllamaClient
from bisheng_client import (
    call_workflow_invoke,
    call_flow_process,
    parse_flow_answer,
    split_to_chunks,
    stop_workflow,
    find_knowledge_id_by_name,
    create_knowledge,
    kb_sync_folder,
)


# 
# --- Bisheng settings (load from CONFIG and allow environment overrides) ---
TAB_ENV_PREFIX = "ENTERPRISE_STANDARD_CHECK"
_BISHENG_CONFIG = CONFIG.get("bisheng", {})
_BISHENG_TAB_CONFIG = _BISHENG_CONFIG.get("tabs", {}).get("enterprise_standard_check", {})


def _bisheng_setting(
    name: str,
    *,
    tab_key: str | None = None,
    config_key: str | None = None,
    default=None,
):
    """Resolve Bisheng settings with tab-specific env and config fallbacks."""

    env_tab_name = f"{TAB_ENV_PREFIX}_{name}"
    env_tab_value = os.getenv(env_tab_name)
    if env_tab_value not in (None, ""):
        return env_tab_value

    env_value = os.getenv(name)
    if env_value not in (None, ""):
        return env_value

    if tab_key:
        tab_value = _BISHENG_TAB_CONFIG.get(tab_key)
        if tab_value not in (None, ""):
            return tab_value

    if config_key:
        config_value = _BISHENG_CONFIG.get(config_key)
        if config_value not in (None, ""):
            return config_value

    return default


def _safe_int(value, fallback: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return fallback


# Base URL and workflow endpoints
BISHENG_BASE_URL = _bisheng_setting("BISHENG_BASE_URL", config_key="base_url", default="http://localhost:3001")
BISHENG_INVOKE_PATH = _bisheng_setting("BISHENG_INVOKE_PATH", config_key="invoke_path", default="/api/v2/workflow/invoke")
BISHENG_STOP_PATH = _bisheng_setting("BISHENG_STOP_PATH", config_key="stop_path", default="/api/v2/workflow/stop")

# Workflow identifiers and API key
BISHENG_WORKFLOW_ID = _bisheng_setting(
    "BISHENG_WORKFLOW_ID",
    tab_key="workflow_id",
    config_key="workflow_id",
    default="",
)
BISHENG_FLOW_ID = _bisheng_setting("BISHENG_FLOW_ID", tab_key="flow_id", default="")
FLOW_INPUT_NODE_ID = _bisheng_setting("FLOW_INPUT_NODE_ID", tab_key="flow_input_node_id", default="")
FLOW_MILVUS_NODE_ID = _bisheng_setting("FLOW_MILVUS_NODE_ID", tab_key="flow_milvus_node_id", default="")
FLOW_ES_NODE_ID = _bisheng_setting("FLOW_ES_NODE_ID", tab_key="flow_es_node_id", default="")
BISHENG_API_KEY = _bisheng_setting("BISHENG_API_KEY", config_key="api_key", default="")

# Chunking and request timeout
BISHENG_MAX_WORDS = _safe_int(_bisheng_setting("BISHENG_MAX_WORDS", config_key="max_words", default=2000), 2000)
BISHENG_TIMEOUT_S = _safe_int(_bisheng_setting("BISHENG_TIMEOUT_S", config_key="timeout_s", default=90), 90)

# Knowledge base settings
# Model ID used when creating a new KB; keep consistent with server defaults
KB_MODEL_ID = 7
# Short slug for this tab; used to synthesize per-user KB names
TAB_SLUG = "enterprise"


def _report_exception(message: str, error: Exception, level: str = "error") -> None:
    """Log exceptions to Streamlit while avoiding silent failures."""

    log_fn = getattr(st, level, None)
    if callable(log_fn):
        log_fn(f"{message}: {error}")
    else:
        st.error(f"{message}: {error}")


def _stream_text(placeholder, text: str, *, chunk_size: int = 30, render_method: str = "text", delay: float | None = None) -> None:
    """Stream text into a Streamlit placeholder in word-sized chunks."""

    words = (text or "").split()
    if not words:
        method = getattr(placeholder, render_method, None)
        if callable(method):
            method("")
        else:
            placeholder.write("")
        return

    method = getattr(placeholder, render_method, None)
    if not callable(method):
        method = placeholder.write

    buffered_words = []
    for start in range(0, len(words), chunk_size):
        buffered_words.append(" ".join(words[start:start + chunk_size]))
        method(" ".join(buffered_words).strip())
        if delay and delay > 0:
            time.sleep(delay)


def _estimate_tokens(text: str) -> int:
    """Rudimentary token estimate: Chinese chars + latin-word count.

    Good enough for correlation studies without server-side tokenizers.
    """
    try:
        cjk = len(re.findall(r"[\u4E00-\u9FFF]", text or ""))
        latin_words = len(re.findall(r"[A-Za-z0-9_]+", text or ""))
        return cjk + latin_words
    except Exception as error:
        _report_exception("令牌估算失败", error, level="warning")
        return max(1, len(text or "") // 2)


def _get_metrics_path(base_out_dir: str, session_id: str) -> str:
    # 每次运行生成一次时间戳文件名并缓存在 session_state
    metrics_dir = os.path.join(base_out_dir, "metrics")
    os.makedirs(metrics_dir, exist_ok=True)
    key = f"metrics_file_{session_id}"
    fname = st.session_state.get(key)
    if not fname:
        from datetime import datetime as _dt
        ts = _dt.now().strftime("%Y%m%d_%H%M%S")
        fname = f"llm_calls_{ts}.csv"
        st.session_state[key] = fname
    return os.path.join(metrics_dir, fname)

def _log_llm_metrics(base_out_dir: str, session_id: str, row: dict):
    try:
        import csv
        path = _get_metrics_path(base_out_dir, session_id)
        exists = os.path.exists(path)
        headers = [
            "ts","engine","model","session_id","file","part","phase",
            "prompt_chars","prompt_tokens","output_chars","output_tokens",
            "duration_ms","success","error"
        ]
        with open(path, 'a', encoding='utf-8-sig', newline='') as f:
            writer = csv.writer(f)
            if not exists:
                writer.writerow(headers)
            writer.writerow([
                row.get("ts"), row.get("engine"), row.get("model"), row.get("session_id"),
                row.get("file"), row.get("part"), row.get("phase"),
                row.get("prompt_chars"), row.get("prompt_tokens"), row.get("output_chars"),
                row.get("output_tokens"), row.get("duration_ms"), row.get("success"), row.get("error")
            ])
    except Exception as error:
        _report_exception("写入LLM指标失败", error, level="warning")


def _persist_compare_outputs(initial_dir: str, name_no_ext: str, prompt_texts: list, full_out_text: str) -> None:
    """Write prompt_{file}.txt and response_{file}.txt under initial_results."""
    try:
        # prompts
        total_parts = len(prompt_texts)
        prompt_out_lines = []
        for idx_p, ptxt in enumerate(prompt_texts, start=1):
            prompt_out_lines.append(f"提示词（第{idx_p}部分，共{total_parts}部分）：")
            prompt_out_lines.append(ptxt)
        prompt_out_text = "\n".join(prompt_out_lines)
        prompt_out_path = os.path.join(initial_dir, f"prompt_{name_no_ext}.txt")
        with open(prompt_out_path, 'w', encoding='utf-8') as pf:
            pf.write(prompt_out_text)
    except Exception as error:
        _report_exception("保存提示词失败", error, level="warning")
    try:
        out_path = os.path.join(initial_dir, f"response_{name_no_ext}.txt")
        with open(out_path, 'w', encoding='utf-8') as outf:
            outf.write(full_out_text)
    except Exception as error:
        _report_exception("保存比对结果失败", error, level="warning")


def _summarize_with_ollama(initial_dir: str, enterprise_out: str, session_id: str, name_no_ext: str, full_out_text: str) -> None:
    """Split long compare text into parts, create prompted files, stream Ollama summarization, log, and write json_* files."""
    # Prepare prompted_response_* files
    try:
        original_name = name_no_ext
        prompt_lines = [
            "你是一个严谨的结构化信息抽取助手。以下内容来自一个基于 RAG 的大语言模型 (RAG-LLM) 系统，",
            "用于将若干技术文件与企业标准进行逐条比对，输出发现的不符合项及其理由。本文件对应的原始技术文件名称为：",
            f"{original_name}。",
            "\n请将随后的全文内容转换为 JSON 数组（list of objects），每个对象包含如下五个键：",
            f"- 技术文件名：\"{original_name}\"",
            "- 技术文件内容：与该条不一致点相关的技术文件条目或段落名称/编号/摘要",
            "- 企业标准：被对比的企业标准条款名称/编号/摘要",
            "- 不一致之处：不符合或存在差异的具体点，尽量具体明确",
            "- 理由：判断不一致的依据与简要解释（保持客观、可追溯）",
            "\n要求：",
            "1) 仅输出严格的 JSON（UTF-8，无注释、无多余文本），键名使用上述中文；",
            "2) 若内容包含多处对比，按条目拆分为多条 JSON 对象；",
            "3) 若某处信息缺失，请以空字符串 \"\" 占位，不要编造；",
            "4) 尽量保留可用于追溯定位的原文线索（如编号、标题、页码等）于相应字段中。",
            "\n下面是需要转换为 JSON 的原始比对输出：\n\n",
        ]
        instr = "".join(prompt_lines)
        text = full_out_text or ""
        # slice by both <think> count and max chars
        max_thinks_per_part = 20
        max_chars_per_part = 8000
        think_indices = []
        needle = "<think>"
        start = 0
        while True:
            pos = text.find(needle, start)
            if pos == -1:
                break
            think_indices.append(pos)
            start = pos + len(needle)
        boundaries = []
        prev = 0
        count_since_prev = 0
        for pos in think_indices:
            count_since_prev += 1
            if count_since_prev >= max_thinks_per_part or (pos - prev) >= max_chars_per_part:
                boundaries.append(pos)
                prev = pos
                count_since_prev = 0
        parts = []
        prev = 0
        for b in boundaries:
            parts.append(text[prev:b])
            prev = b
        parts.append(text[prev:])
        for idx_part, part in enumerate(parts, start=1):
            cleaned_part = re.sub(r"<think>[\s\S]*?</think>", "", part)
            prompted_path = os.path.join(initial_dir, f"prompted_response_{name_no_ext}_pt{idx_part}.txt")
            with open(prompted_path, 'w', encoding='utf-8') as pf:
                pf.write(instr)
                pf.write(cleaned_part)
    except Exception as error:
        _report_exception("生成Ollama汇总提示失败", error, level="warning")

    # Summarize each prompted part with Ollama
    try:
        # collect parts
        part_files = []
        try:
            for fn in os.listdir(initial_dir):
                if fn.startswith(f"prompted_response_{name_no_ext}_pt") and fn.endswith('.txt'):
                    part_files.append(fn)
        except Exception as error:
            _report_exception("读取Ollama提示列表失败", error, level="warning")
            part_files = []
        part_files.sort(key=lambda x: x.lower())
        total_parts = len(part_files)
        host = resolve_ollama_host("ollama_9")
        ollama_client = OllamaClient(host=host)
        model = CONFIG["llm"].get("ollama_model")
        temperature = 0.7
        top_p = 0.9
        top_k = 40
        repeat_penalty = 1.1
        num_ctx = 40001
        num_thread = 4
        for part_idx, pfname in enumerate(part_files, start=1):
            p_path = os.path.join(initial_dir, pfname)
            try:
                with open(p_path, 'r', encoding='utf-8') as fpr:
                    prompt_text_all = fpr.read()
            except Exception as error:
                _report_exception(f"读取汇总提示失败({pfname})", error, level="warning")
                prompt_text_all = ""

            base = pfname[len("prompted_response_"):]
            json_name = f"json_{base}"
            json_path = os.path.join(initial_dir, json_name)
            existing_json_text = ""
            if os.path.isfile(json_path):
                try:
                    with open(json_path, 'r', encoding='utf-8') as jf:
                        existing_json_text = jf.read()
                except Exception as error:
                    _report_exception(f"读取已存在的JSON汇总失败({json_name})", error, level="warning")
                    existing_json_text = ""

            # columns
            col_prompt2, col_resp2 = st.columns([1, 1])
            with col_prompt2:
                st.markdown(f"生成汇总表格提示词（第{part_idx}部分，共{total_parts}部分）")
                prompt_container2 = st.container(height=400)
                with prompt_container2:
                    with st.chat_message("user"):
                        ph2 = st.empty()
                        _stream_text(ph2, prompt_text_all, render_method="text")
                st.chat_input(placeholder="", disabled=True, key=f"enterprise_prompted_prompt_{session_id}_{pfname}")
            with col_resp2:
                st.markdown(f"生成汇总表格结果（第{part_idx}部分，共{total_parts}部分）")
                resp_container2 = st.container(height=400)
                with resp_container2:
                    with st.chat_message("assistant"):
                        ph_resp2 = st.empty()
                        if (existing_json_text or "").strip():
                            _stream_text(ph_resp2, existing_json_text, render_method="write")
                            st.caption("已载入之前生成的JSON结果，无需重新调用模型。")
                        else:
                            response_text = ""
                            start_ts = time.time()
                            last_stats = None
                            error_msg = ""
                            chunks_seen = 0
                            try:
                                for chunk in ollama_client.chat(
                                    model=model,
                                    messages=[{"role": "user", "content": prompt_text_all}],
                                    stream=True,
                                    options={
                                        "temperature": temperature,
                                        "top_p": top_p,
                                        "top_k": top_k,
                                        "repeat_penalty": repeat_penalty,
                                        "num_ctx": num_ctx,
                                        "num_thread": num_thread,
                                    },
                                ):
                                    chunks_seen += 1
                                    new_text = chunk.get('message', {}).get('content', '')
                                    response_text += new_text
                                    ph_resp2.write(response_text)
                                    last_stats = chunk.get('eval_info') or chunk.get('stats') or last_stats
                            except Exception as e:
                                error_msg = str(e)[:300]
                            finally:
                                from datetime import datetime as _dt

                                dur_ms = int((time.time() - start_ts) * 1000)
                                prompt_tokens = (last_stats or {}).get('prompt_eval_count')
                                output_tokens = (last_stats or {}).get('eval_count')
                                if not error_msg and chunks_seen == 0:
                                    error_msg = "no_stream_chunks"
                                _log_llm_metrics(
                                    enterprise_out,
                                    session_id,
                                    {
                                        "ts": _dt.now().isoformat(timespec="seconds"),
                                        "engine": "ollama",
                                        "model": model,
                                        "session_id": "",
                                        "file": name_no_ext,
                                        "part": part_idx,
                                        "phase": "summarize",
                                        "prompt_chars": len(prompt_text_all or ""),
                                        "prompt_tokens": prompt_tokens if isinstance(prompt_tokens, int) else _estimate_tokens(prompt_text_all or ""),
                                        "output_chars": len(response_text or ""),
                                        "output_tokens": output_tokens if isinstance(output_tokens, int) else _estimate_tokens(response_text or ""),
                                        "duration_ms": dur_ms,
                                        "success": 1 if (response_text or "").strip() else 0,
                                        "error": error_msg,
                                    },
                                )
                                # save json
                                try:
                                    import tempfile as _tmp, shutil as _sh

                                    with _tmp.NamedTemporaryFile('w', delete=False, encoding='utf-8', dir=initial_dir) as jf_tmp:
                                        jf_tmp.write(response_text)
                                        tmp_name = jf_tmp.name
                                    _sh.move(tmp_name, json_path)
                                except Exception as error:
                                    _report_exception("写入JSON汇总失败", error, level="warning")
                        response_text = ""
                        start_ts = time.time()
                        last_stats = None
                        error_msg = ""
                        chunks_seen = 0
                        try:
                            for chunk in ollama_client.chat(
                                model=model,
                                messages=[{"role": "user", "content": prompt_text_all}],
                                stream=True,
                                options={ "temperature": temperature, "top_p": top_p, "top_k": top_k,
                                          "repeat_penalty": repeat_penalty, "num_ctx": num_ctx, "num_thread": num_thread }
                            ):
                                chunks_seen += 1
                                new_text = chunk.get('message', {}).get('content', '')
                                response_text += new_text
                                ph_resp2.write(response_text)
                                last_stats = chunk.get('eval_info') or chunk.get('stats') or last_stats
                        except Exception as e:
                            error_msg = str(e)[:300]
                        finally:
                            from datetime import datetime as _dt
                            dur_ms = int((time.time() - start_ts) * 1000)
                            prompt_tokens = (last_stats or {}).get('prompt_eval_count')
                            output_tokens = (last_stats or {}).get('eval_count')
                            if not error_msg and chunks_seen == 0:
                                error_msg = "no_stream_chunks"
                            _log_llm_metrics(
                                enterprise_out, session_id,
                                {
                                    "ts": _dt.now().isoformat(timespec="seconds"),
                                    "engine": "ollama",
                                    "model": model,
                                    "session_id": "",
                                    "file": name_no_ext,
                                    "part": part_idx,
                                    "phase": "summarize",
                                    "prompt_chars": len(prompt_text_all or ""),
                                    "prompt_tokens": prompt_tokens if isinstance(prompt_tokens, int) else _estimate_tokens(prompt_text_all or ""),
                                    "output_chars": len(response_text or ""),
                                    "output_tokens": output_tokens if isinstance(output_tokens, int) else _estimate_tokens(response_text or ""),
                                    "duration_ms": dur_ms,
                                    "success": 1 if (response_text or "").strip() else 0,
                                    "error": error_msg
                                }
                            )
                            # save json
                            try:
                                base = pfname[len("prompted_response_"):]
                                json_name = f"json_{base}"
                                json_path = os.path.join(initial_dir, json_name)
                                with open(json_path, 'w', encoding='utf-8') as jf:
                                    jf.write(response_text)
                            except Exception as error:
                                _report_exception("写入JSON汇总失败", error, level="warning")
    except Exception as error:
        _report_exception("调用Ollama生成汇总失败", error)


def _aggregate_outputs(initial_dir: str, enterprise_out: str, session_id: str) -> None:
    """Aggregate json_* into CSV/XLSX, and generate Word document from prompt_/response_."""
    final_dir = os.path.join(enterprise_out, 'final_results')
    os.makedirs(final_dir, exist_ok=True)
    # collect json
    json_files = []
    try:
        for fn in os.listdir(initial_dir):
            if fn.startswith('json_') and fn.lower().endswith('.txt'):
                json_files.append(os.path.join(initial_dir, fn))
    except Exception as error:
        _report_exception("读取初始结果目录失败", error, level="warning")
        json_files = []
    columns = ["技术文件名", "技术文件内容", "企业标准", "不一致之处", "理由"]
    rows = []
    import csv
    try:
        import pandas as pd  # type: ignore
    except ImportError as error:
        pd = None  # type: ignore[assignment]
        st.info(f"未安装 pandas，Excel 导出将被跳过：{error}")
    def _strip_code_fences(s: str) -> str:
        s = (s or "").strip()
        if s.startswith("```") and s.endswith("```"):
            s = s[3:-3].strip()
            if s.lower().startswith("json"):
                s = s[4:].strip()
        return s
    def _extract_json_text(text: str) -> str:
        s = (text or "").strip().lstrip("\ufeff")
        s = _strip_code_fences(s)
        try:
            json.loads(s)
            return s
        except json.JSONDecodeError:
            pass
        try:
            start = s.find("[")
            end = s.rfind("]")
            if start != -1 and end != -1 and end > start:
                candidate = s[start:end+1]
                try:
                    json.loads(candidate)
                    return candidate
                except json.JSONDecodeError:
                    merged = re.sub(r"\]\s*,?\s*\[", ",", candidate)
                    json.loads(merged)
                    return merged
        except json.JSONDecodeError:
            pass
        try:
            start = s.find("{")
            end = s.rfind("}")
            if start != -1 and end != -1 and end > start:
                candidate = s[start:end+1]
                json.loads(candidate)
                return candidate
        except json.JSONDecodeError:
            pass
        return s
    for jf in json_files:
        try:
            with open(jf, 'r', encoding='utf-8') as f:
                text = f.read()
        except Exception as error:
            _report_exception(f"读取JSON结果失败({os.path.basename(jf)})", error, level="warning")
            text = ""
        js = _extract_json_text(text)
        try:
            parsed = json.loads(js)
        except json.JSONDecodeError:
            parsed = None
        items = []
        if isinstance(parsed, list):
            items = [x for x in parsed if isinstance(x, dict)]
        elif isinstance(parsed, dict):
            vals = list(parsed.values())
            for v in vals:
                if isinstance(v, list) and all(isinstance(i, dict) for i in v):
                    items = v
                    break
            if not items:
                items = [parsed]
        base_name = os.path.basename(jf)
        orig_name = base_name[5:] if base_name.startswith('json_') else base_name
        try:
            orig_name = re.sub(r"_pt\d+\.txt$", "", orig_name)
        except Exception:
            pass
        for obj in items:
            row = [
                orig_name,
                str(obj.get("技术文件内容", "")),
                str(obj.get("企业标准", "")),
                str(obj.get("不一致之处", "")),
                str(obj.get("理由", "")),
            ]
            rows.append(row)
    from datetime import datetime as _dt
    ts = _dt.now().strftime('%Y%m%d_%H%M%S')
    csv_path = os.path.join(final_dir, f"企标检查结果_{ts}.csv")
    with open(csv_path, 'w', encoding='utf-8-sig', newline='') as cf:
        writer = csv.writer(cf)
        writer.writerow(columns)
        for r in rows:
            writer.writerow(r)
    xlsx_path = None
    if pd is not None:
        try:
            df = pd.DataFrame(rows, columns=columns)
            xlsx_path = os.path.join(final_dir, f"企标检查结果_{ts}.xlsx")
            df.to_excel(xlsx_path, index=False, engine='openpyxl')
        except Exception as error:
            _report_exception("生成Excel结果失败", error, level="warning")
            xlsx_path = None
    try:
        with open(csv_path, 'rb') as fcsv:
            st.download_button(label="下载CSV结果", data=fcsv.read(), file_name=os.path.basename(csv_path), mime='text/csv', key=f"download_csv_{session_id}")
        if xlsx_path and os.path.exists(xlsx_path):
            with open(xlsx_path, 'rb') as fxlsx:
                st.download_button(label="下载Excel结果", data=fxlsx.read(), file_name=os.path.basename(xlsx_path), mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', key=f"download_xlsx_{session_id}")
    except Exception as error:
        _report_exception("生成下载链接失败", error, level="warning")
    try:
        try:
            from docx import Document
        except Exception:
            Document = None
        if Document is not None:
            pairs = []
            prompt_files = [f for f in os.listdir(initial_dir) if f.startswith('prompt_') and f.lower().endswith('.txt')]
            for pf in prompt_files:
                base = pf[len('prompt_'):-4]
                rf = os.path.join(initial_dir, f"response_{base}.txt")
                if os.path.isfile(rf):
                    pairs.append((os.path.join(initial_dir, pf), rf, base))
            pairs.sort(key=lambda x: x[2].lower())
            if pairs:
                doc = Document()
                try:
                    styles = doc.styles
                    styles['Normal'].font.name = '宋体'
                    styles['Heading 1'].font.name = '宋体'
                    styles['Heading 2'].font.name = '宋体'
                except Exception:
                    pass
                doc.add_heading("企标检查全部分析过程", level=0)
                doc.add_heading("目录", level=1)
                for _, _, base in pairs:
                    p = doc.add_paragraph()
                    p.add_run(f"{base}")
                for p_path, r_path, base in pairs:
                    doc.add_heading(f"以下是《{base}》根据企业标准检查的分析过程：", level=1)
                    try:
                        with open(p_path, 'r', encoding='utf-8') as f:
                            ptext = f.read()
                    except Exception as error:
                        _report_exception(f"读取提示词文件失败({os.path.basename(p_path)})", error, level="warning")
                        ptext = ""
                    def _to_plain_table(txt: str) -> str:
                        try:
                            import re as _re
                            rows = []
                            for m in _re.finditer(r"<tr[\s\S]*?>([\s\S]*?)</tr>", txt, flags=_re.IGNORECASE):
                                row_html = m.group(1)
                                cells = _re.findall(r"<(?:td|th)[^>]*>([\s\S]*?)</(?:td|th)>", row_html, flags=_re.IGNORECASE)
                                cells_clean = [ _re.sub(r"<[^>]+>", "", c).strip() for c in cells ]
                                if cells_clean:
                                    rows.append("\t".join(cells_clean))
                            if rows:
                                return "\n".join(rows)
                            return _re.sub(r"<[^>]+>", " ", txt)
                        except Exception:
                            return txt
                    for para in (ptext or "").splitlines():
                        line = para.strip()
                        if not line:
                            continue
                        if re.match(r"^提示词（第\d+部分，共\d+部分）：", line):
                            doc.add_heading(line, level=2)
                        else:
                            doc.add_paragraph(_to_plain_table(line))
                    try:
                        with open(r_path, 'r', encoding='utf-8') as f:
                            rtext = f.read()
                    except Exception as error:
                        _report_exception(f"读取比对结果失败({os.path.basename(r_path)})", error, level="warning")
                        rtext = ""
                    for para in (rtext or "").splitlines():
                        line = para.strip()
                        if not line:
                            continue
                        doc.add_paragraph(_to_plain_table(line))
                from datetime import datetime as _dt
                ts_doc = _dt.now().strftime('%Y%m%d_%H%M%S')
                doc_path = os.path.join(final_dir, f"企标检查分析过程_{ts_doc}.docx")
                doc.save(doc_path)
                try:
                    with open(doc_path, 'rb') as fdoc:
                        st.download_button(label="下载分析过程Word", data=fdoc.read(), file_name=os.path.basename(doc_path), mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', key=f"download_docx_{session_id}")
                except Exception:
                    pass
        else:
            st.info("未安装 python-docx，暂无法生成Word文档。请安装 python-docx 后重试。")
    except Exception as e:
        st.error(f"生成分析过程Word失败：{e}")

def _list_pdfs(folder: str):
    """Return absolute paths for all PDF files in a folder (non-recursive)."""
    try:
        return [
            os.path.join(folder, f)
            for f in os.listdir(folder)
            if os.path.isfile(os.path.join(folder, f)) and f.lower().endswith('.pdf')
        ]
    except Exception:
        return []


def _mineru_parse_pdf(pdf_path: str) -> bytes:
    """Call MinerU API to parse a single PDF and return ZIP bytes on success.

    Raises an exception on failure.
    """
    api_url = "http://10.31.60.127:8000/file_parse"
    data = {
        'backend': 'vlm-sglang-engine',
        'response_format_zip': 'true',
        # Enable richer outputs; we will primarily consume the .md text for now
        'formula_enable': 'true',
        'table_enable': 'true',
        'return_images': 'false',
        'return_middle_json': 'true',
        'return_model_output': 'false',
        'return_content_list': 'true',
    }
    with open(pdf_path, 'rb') as f:
        files = {'files': (os.path.basename(pdf_path), f, 'application/pdf')}
        resp = requests.post(api_url, data=data, files=files, timeout=300)
        if resp.status_code != 200:
            raise RuntimeError(f"MinerU API error {resp.status_code}: {resp.text[:200]}")
        return resp.content


def _zip_to_txts(zip_bytes: bytes, target_txt_path: str) -> bool:
    """Extract first .md file from ZIP bytes and save as plain text (.txt).

    Returns True if a .txt was written, False otherwise.
    """
    # MinerU returns a ZIP archive for each PDF containing: a Markdown file (extracted
    # plain text content), JSONs (structured intermediates), and optionally images.
    # For now we only need plain text for LLM prompts, so we take the first .md file
    # and write it out as a .txt. The images are intentionally ignored here, but they
    # are valuable for future RAG/Q&A over figures or diagrams. We will revisit image
    # handling later to index them alongside text for multimodal retrieval.
    bio = io.BytesIO(zip_bytes)
    try:
        with zipfile.ZipFile(bio) as zf:
            # Prefer top-level or nested .md
            md_members = [n for n in zf.namelist() if n.lower().endswith('.md')]
            if not md_members:
                return False
            # Use the first .md
            name = md_members[0]
            content = zf.read(name)
            # Ensure output directory exists
            os.makedirs(os.path.dirname(target_txt_path), exist_ok=True)
            with open(target_txt_path, 'wb') as out_f:
                out_f.write(content)
            return True
    except zipfile.BadZipFile:
        return False

def _insert_source_markers(text: str, source_label: str, line_interval: int = 80) -> str:
    """Insert unobtrusive source markers so small retrieved fragments retain provenance.

    Strategy:
    - Add a file-level header at top: 【来源文件: <name>】
    - If Markdown headings (#/##) are present, add a marker right after each H1/H2.
    - Otherwise, add a marker every N non-empty lines (default 80).

    Idempotent: if a marker containing this source_label already exists, return original text.
    """
    marker = f"【来源文件: {source_label}】"
    if source_label and marker in text:
        return text
    lines = text.splitlines()
    has_md_heading = any(re.match(r'^\s{0,3}#{1,2}\s+\S', ln) for ln in lines[:500])

    annotated_lines = []
    # Always place a header at the very top
    annotated_lines.append(marker)
    annotated_lines.append("")

    if has_md_heading:
        for ln in lines:
            annotated_lines.append(ln)
            if re.match(r'^\s{0,3}#{1,2}\s+\S', ln):
                annotated_lines.append(marker)
    else:
        non_empty_count = 0
        for ln in lines:
            annotated_lines.append(ln)
            if ln.strip():
                non_empty_count += 1
                if non_empty_count % max(10, int(line_interval)) == 0:
                    annotated_lines.append(marker)

    return "\n".join(annotated_lines)


def _annotate_txt_file_inplace(file_path: str, source_label: str, line_interval: int = 80) -> bool:
    """Open a .txt file and inject source markers in-place. Returns True if updated."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            original = f.read()
        annotated = _insert_source_markers(original, source_label, line_interval=line_interval)
        if annotated == original:
            return False
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(annotated)
        return True
    except Exception:
        return False

def _process_pdf_folder(input_dir: str, output_dir: str, progress_area, annotate_sources: bool = False):
    """Process all PDFs in input_dir via MinerU and write .txts into output_dir."""
    pdf_paths = _list_pdfs(input_dir)
    if not pdf_paths:
        return []
    created = []
    for pdf_path in pdf_paths:
        orig_name = os.path.basename(pdf_path)
        # Preserve original extension in output filename, e.g., name.pdf.txt
        out_txt = os.path.join(output_dir, f"{orig_name}.txt")
        try:
            # Skip if parsed file already exists and is non-empty
            if os.path.exists(out_txt) and os.path.getsize(out_txt) > 0:
                progress_area.info(f"已存在（跳过）: {os.path.basename(out_txt)}")
                continue
            progress_area.write(f"解析: {os.path.basename(pdf_path)} …")
            zip_bytes = _mineru_parse_pdf(pdf_path)
            ok = _zip_to_txts(zip_bytes, out_txt)
            if ok:
                # Inject source markers so retrieved snippets carry provenance
                if annotate_sources:
                    _annotate_txt_file_inplace(out_txt, orig_name)
                created.append(out_txt)
            else:
                progress_area.warning(f"未发现可用的 .md 内容，跳过: {os.path.basename(pdf_path)}")
        except Exception as e:
            progress_area.error(f"失败: {os.path.basename(pdf_path)} → {e}")
    return created


def _list_word_ppt(folder: str):
    """Return absolute paths for .doc, .docx, .ppt, .pptx in a folder (non-recursive)."""
    try:
        return [
            os.path.join(folder, f)
            for f in os.listdir(folder)
            if os.path.isfile(os.path.join(folder, f)) and os.path.splitext(f)[1].lower() in {'.doc', '.docx', '.ppt', '.pptx'}
        ]
    except Exception:
        return []


def _unstructured_partition_to_txt(file_path: str, target_txt_path: str) -> bool:
    """Send a single Word/PPT file to Unstructured API and write plain text (.txt).

    The Unstructured server is expected at 10.31.60.11 running the API. We call the
    "general" endpoint and extract text fields. Table-like data, if present, are
    flattened with tab separators for readability in plain text.

    Future plan (RAG-focused tables):
    - Keep Unstructured for narrative text/structure.
    - Extract tables directly from the original DOCX/PPTX using python-docx/python-pptx.
    - Convert those tables to TSV (one row per line, cells separated by a single tab).
    - Replace or insert TSV blocks into the final output .txt in place of flattened table text.
    This will improve recall/precision on numeric lookups. Not implemented yet; will
    be added later when schedule allows.
    """
    # Resolve API URL: env var first, then CONFIG.services.unstructured_api_url
    api_url = os.getenv('UNSTRUCTURED_API_URL') or CONFIG.get('services', {}).get('unstructured_api_url') or 'http://10.31.60.11:8000/general/v0/general'
    try:
        with open(file_path, 'rb') as f:
            files = {'files': (os.path.basename(file_path), f)}
            # RAG-optimized defaults: structured tables, auto strategy, Chinese+English OCR support
            form = {
                "strategy": "auto",
                "ocr_languages": "chi_sim,eng",
                "infer_table_structure": "true",
            }
            resp = requests.post(api_url, files=files, data=form, timeout=300)
            if resp.status_code != 200:
                raise RuntimeError(f"Unstructured API {resp.status_code}: {resp.text[:200]}")
            data = resp.json()
            # data is expected to be a list of elements; each may have 'text' or table-like content
            lines = []
            if isinstance(data, list):
                for el in data:
                    # Prefer 'text'
                    text = None
                    if isinstance(el, dict):
                        # Common key is 'text'
                        text = el.get('text')
                        # Some table extractions might be under 'data' (list of rows)
                        if not text and isinstance(el.get('data'), list):
                            for row in el['data']:
                                if isinstance(row, list):
                                    lines.append('\t'.join(str(c) for c in row))
                            # Continue to next element after adding table rows
                            continue
                    if isinstance(text, str) and text.strip():
                        lines.append(text.strip())
            # Write as UTF-8 plain text
            os.makedirs(os.path.dirname(target_txt_path), exist_ok=True)
            with open(target_txt_path, 'w', encoding='utf-8') as out_f:
                out_f.write('\n'.join(lines))
            return True
    except Exception as e:
        # Surface errors to caller via return False; logging via progress UI
        return False


def _process_word_ppt_folder(input_dir: str, output_dir: str, progress_area, annotate_sources: bool = False):
    """Process .doc/.docx/.ppt/.pptx via Unstructured API and write .txts."""
    paths = _list_word_ppt(input_dir)
    if not paths:
        return []
    created = []
    for p in paths:
        orig_name = os.path.basename(p)
        # Preserve original extension in output filename, e.g., name.docx.txt / name.ppt.txt
        out_txt = os.path.join(output_dir, f"{orig_name}.txt")
        try:
            # Skip if parsed file already exists and is non-empty
            if os.path.exists(out_txt) and os.path.getsize(out_txt) > 0:
                progress_area.info(f"已存在（跳过）: {os.path.basename(out_txt)}")
                continue
            progress_area.write(f"解析(Word/PPT): {os.path.basename(p)} …")
            ok = _unstructured_partition_to_txt(p, out_txt)
            if ok:
                # Inject source markers
                if annotate_sources:
                    _annotate_txt_file_inplace(out_txt, orig_name)
                created.append(out_txt)
            else:
                progress_area.warning(f"未能从文件中生成文本，跳过: {os.path.basename(p)}")
        except Exception as e:
            progress_area.error(f"失败: {os.path.basename(p)} → {e}")
    return created


def _list_excels(folder: str):
    """Return absolute paths for .xls/.xlsx/.xlsm in a folder (non-recursive)."""
    try:
        return [
            os.path.join(folder, f)
            for f in os.listdir(folder)
            if os.path.isfile(os.path.join(folder, f)) and os.path.splitext(f)[1].lower() in {'.xls', '.xlsx', '.xlsm'}
        ]
    except Exception:
        return []


def _sanitize_sheet_name(name: str) -> str:
    """Sanitize sheet names for filenames: keep readable, remove path-forbidden chars."""
    bad = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
    for ch in bad:
        name = name.replace(ch, '_')
    return '_'.join(name.strip().split())[:80] or 'Sheet'


def _process_excel_folder(input_dir: str, output_dir: str, progress_area, annotate_sources: bool = False):
    """Convert each Excel sheet to CSV text and save as <file>_SHEET_<sheet>.txt.

    Note: We intentionally save CSV content with a .txt extension for uniform LLM
    consumption. This is technically fine: the content is plain text CSV and the
    file extension does not affect parsing for our use case.
    """
    paths = _list_excels(input_dir)
    if not paths:
        return []
    created = []
    try:
        import pandas as pd  # type: ignore
    except ImportError as error:
        progress_area.warning(f"未安装 pandas，无法处理 Excel：{error}")
        return []
    for excel_path in paths:
        orig_name = os.path.basename(excel_path)  # keep extension in base name per spec
        try:
            xls = pd.ExcelFile(excel_path)
            for sheet in xls.sheet_names:
                safe_sheet = _sanitize_sheet_name(sheet)
                out_txt = os.path.join(output_dir, f"{orig_name}_SHEET_{safe_sheet}.txt")
                # Skip if exists and non-empty
                if os.path.exists(out_txt) and os.path.getsize(out_txt) > 0:
                    progress_area.info(f"已存在（跳过）: {os.path.basename(out_txt)}")
                    continue
                progress_area.write(f"转换(Excel→CSV): {orig_name} / {sheet} …")
                df = xls.parse(sheet)
                # Write CSV content into .txt
                df.to_csv(out_txt, index=False, encoding='utf-8')
                # Inject source markers including sheet context
                if annotate_sources:
                    _annotate_txt_file_inplace(out_txt, f"{orig_name} / {sheet}")
                created.append(out_txt)
        except Exception as e:
            progress_area.error(f"失败: {orig_name} → {e}")
    return created


def _process_textlike_folder(input_dir: str, output_dir: str, progress_area):
    """Copy text-like files (csv, tsv, md, txt, json, yaml, yml, log, ini, cfg, rst)
    to output_dir with .txt extension. Skip if the target exists and is non-empty.
    """
    try:
        if not os.path.isdir(input_dir):
            return []
        exts = {
            '.txt', '.md', '.csv', '.tsv', '.json', '.yaml', '.yml', '.log', '.ini', '.cfg', '.rst'
        }
        os.makedirs(output_dir, exist_ok=True)
        written = []
        for name in os.listdir(input_dir):
            spath = os.path.join(input_dir, name)
            if not os.path.isfile(spath):
                continue
            ext = os.path.splitext(name)[1].lower()
            if ext not in exts:
                continue
            base = os.path.splitext(name)[0]
            dst = os.path.join(output_dir, f"{base}.txt")
            try:
                if os.path.exists(dst) and os.path.getsize(dst) > 0:
                    progress_area.info(f"已存在（跳过）: {os.path.basename(dst)}")
                    continue
                with open(spath, 'r', encoding='utf-8', errors='ignore') as fr:
                    content = fr.read()
                with open(dst, 'w', encoding='utf-8') as fw:
                    fw.write(content)
                written.append(dst)
            except Exception as e:
                progress_area.warning(f"复制失败: {name} → {e}")
        return written
    except Exception:
        return []


def _process_archives(input_dir: str, output_dir: str, progress_area) -> int:
    """Extract archives (.zip, optionally .7z/.rar if libs available) and process their contents
    into text outputs under output_dir. Returns number of archives processed.

    Strategy:
    - For each archive in input_dir, extract into a temp folder under output_dir.
    - For each extracted subfolder, run the existing processors on that folder:
      PDFs → MinerU; Word/PPT → Unstructured; Excel → CSV→.txt; text-like → copy.
    - Clean up extracted temp if possible.
    """
    try:
        if not os.path.isdir(input_dir):
            return 0
        import shutil
        import tempfile
        processed = 0
        # Optional handlers
        try:
            import py7zr  # type: ignore
        except Exception:
            py7zr = None  # type: ignore
        try:
            import rarfile  # type: ignore
        except Exception:
            rarfile = None  # type: ignore

        for name in os.listdir(input_dir):
            spath = os.path.join(input_dir, name)
            if not os.path.isfile(spath):
                continue
            ext = os.path.splitext(name)[1].lower()
            if ext not in {'.zip', '.7z', '.rar'}:
                continue
            # Create temp extraction root
            tmp_root = tempfile.mkdtemp(prefix="extract_", dir=output_dir)
            ok = False
            try:
                if ext == '.zip':
                    try:
                        with zipfile.ZipFile(spath) as zf:
                            zf.extractall(tmp_root)
                        ok = True
                    except Exception as e:
                        progress_area.warning(f"解压失败: {name} → {e}")
                elif ext == '.7z' and py7zr is not None:
                    try:
                        with py7zr.SevenZipFile(spath, mode='r') as z:
                            z.extractall(path=tmp_root)
                        ok = True
                    except Exception as e:
                        progress_area.warning(f"解压7z失败: {name} → {e}")
                elif ext == '.rar' and rarfile is not None:
                    try:
                        rf = rarfile.RarFile(spath)
                        rf.extractall(tmp_root)
                        rf.close()
                        ok = True
                    except Exception as e:
                        progress_area.warning(f"解压rar失败: {name} → {e}")
                else:
                    # Unsupported archive (missing libs)
                    progress_area.info(f"跳过未支持的压缩包: {name}")

                if ok:
                    # Walk extracted tree and process each folder with existing processors
                    for root, dirs, files in os.walk(tmp_root):
                        # Run per-folder handlers (they skip existing outputs)
                        _process_pdf_folder(root, output_dir, progress_area, annotate_sources=True)
                        _process_word_ppt_folder(root, output_dir, progress_area, annotate_sources=True)
                        _process_excel_folder(root, output_dir, progress_area, annotate_sources=True)
                        _process_textlike_folder(root, output_dir, progress_area)
                    processed += 1
            finally:
                # Cleanup temp extraction root
                try:
                    shutil.rmtree(tmp_root, ignore_errors=True)
                except Exception:
                    pass
        return processed
    except Exception:
        return 0

def _cleanup_orphan_txts(source_dir: str, output_dir: str, progress_area=None) -> int:
    """Delete .txt files in output_dir that do not correspond to files currently in source_dir.

    Rules:
    - For PDF/Word/PPT sources: keep exact name match "<orig_name>.txt" where <orig_name> includes the extension
      (e.g., foo.pdf -> foo.pdf.txt, bar.docx -> bar.docx.txt).
    - For Excel sources: keep any file starting with "<orig_name>_SHEET_" (multiple per workbook).
    """
    try:
        if not os.path.isdir(output_dir):
            return 0
        # Build allowlists from current source files
        keep_exact = set()
        keep_prefixes = []
        try:
            for fname in os.listdir(source_dir or "."):
                spath = os.path.join(source_dir, fname)
                if not os.path.isfile(spath):
                    continue
                ext = os.path.splitext(fname)[1].lower()
                if ext in {'.pdf', '.doc', '.docx', '.ppt', '.pptx'}:
                    keep_exact.add((fname + '.txt').lower())
                elif ext in {'.xls', '.xlsx', '.xlsm'}:
                    keep_prefixes.append((fname + '_SHEET_').lower())
        except Exception:
            # If source_dir not accessible, do not delete anything
            return 0

        deleted_count = 0
        for oname in os.listdir(output_dir):
            opath = os.path.join(output_dir, oname)
            if not os.path.isfile(opath):
                continue
            name_lower = oname.lower()
            if not name_lower.endswith('.txt'):
                continue
            keep = name_lower in keep_exact or any(name_lower.startswith(pfx) for pfx in keep_prefixes)
            if not keep:
                try:
                    os.remove(opath)
                    deleted_count += 1
                    if progress_area is not None:
                        progress_area.info(f"清理无关文本: {oname}")
                except Exception:
                    pass
        return deleted_count
    except Exception:
        return 0


def render_enterprise_standard_check_tab(session_id):
    # Handle None session_id (user not logged in)
    if session_id is None:
        st.warning("请先登录以使用此功能。")
        return

    st.subheader("🏢 企业标准检查")

    # No CSS width overrides; rely on Streamlit columns like special symbols tab
    # Ensure enterprise directories and a generated output root exist
    base_dirs = {
        "generated": str(CONFIG["directories"]["generated_files"]),
    }
    session_dirs = ensure_session_dirs(base_dirs, session_id)
    standards_dir = session_dirs.get("enterprise_standards")
    examined_dir = session_dirs.get("enterprise_examined")
    generated_session_dir = session_dirs.get("generated")
    enterprise_out_root = os.path.join(generated_session_dir, "enterprise_standard_check")
    standards_txt_dir = os.path.join(enterprise_out_root, "standards_txt")
    examined_txt_dir = os.path.join(enterprise_out_root, "examined_txt")
    os.makedirs(standards_txt_dir, exist_ok=True)
    os.makedirs(examined_txt_dir, exist_ok=True)

    # Layout: right column for info, left for main content
    col_main, col_info = st.columns([2, 1])

    with col_info:
        # Right column intentionally limited to file manager and utilities only
        # File manager utilities (mirroring completeness tab behavior)
        def get_file_list(folder):
            if not folder or not os.path.exists(folder):
                return []
            files = []
            for f in os.listdir(folder):
                file_path = os.path.join(folder, f)
                if os.path.isfile(file_path):
                    stat = os.stat(file_path)
                    files.append({
                        'name': f,
                        'size': stat.st_size,
                        'modified': stat.st_mtime,
                        'path': file_path
                    })
            # Sort by name then modified time for stability
            return sorted(files, key=lambda x: (x['name'].lower(), x['modified']))

        def format_file_size(size_bytes):
            if size_bytes == 0:
                return "0 B"
            size_names = ["B", "KB", "MB", "GB"]
            i = 0
            while size_bytes >= 1024 and i < len(size_names) - 1:
                size_bytes /= 1024.0
                i += 1
            return f"{size_bytes:.1f} {size_names[i]}"

        def format_timestamp(timestamp):
            from datetime import datetime
            return datetime.fromtimestamp(timestamp).strftime('%Y-%m-%d %H:%M')

        def truncate_filename(filename, max_length=40):
            if len(filename) <= max_length:
                return filename
            name, ext = os.path.splitext(filename)
            available_length = max_length - len(ext) - 3
            if available_length <= 0:
                return filename[:max_length-3] + "..."
            truncated_name = name[:available_length] + "..."
            return truncated_name + ext

        # Clear buttons
        col_clear1, col_clear2, col_clear3 = st.columns(3)
        with col_clear1:
            if st.button("🗑️ 清空企业标准文件", key=f"clear_enterprise_std_{session_id}"):
                try:
                    for file in os.listdir(standards_dir):
                        file_path = os.path.join(standards_dir, file)
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                    st.success("已清空企业标准文件")
                except Exception as e:
                    st.error(f"清空失败: {e}")
        with col_clear2:
            if st.button("🗑️ 清空待检查文件", key=f"clear_enterprise_exam_{session_id}"):
                try:
                    for file in os.listdir(examined_dir):
                        file_path = os.path.join(examined_dir, file)
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                    st.success("已清空待检查文件")
                except Exception as e:
                    st.error(f"清空失败: {e}")
        with col_clear3:
            if st.button("🗑️ 清空分析结果", key=f"clear_enterprise_results_{session_id}"):
                try:
                    final_dir_path = os.path.join(enterprise_out_root, 'final_results')
                    deleted_count = 0
                    if os.path.isdir(final_dir_path):
                        for fname in os.listdir(final_dir_path):
                            fpath = os.path.join(final_dir_path, fname)
                            if os.path.isfile(fpath):
                                os.remove(fpath)
                                deleted_count += 1
                    st.success(f"已清空分析结果（{deleted_count} 个文件）")
                except Exception as e:
                    st.error(f"清空失败: {e}")

        # File lists in tabs (fixed order)
        tab_std, tab_exam, tab_results = st.tabs(["企业标准文件", "待检查文件", "分析结果"])
        with tab_std:
            std_files = get_file_list(standards_dir)
            if std_files:
                for file_info in std_files:
                    display_name = truncate_filename(file_info['name'])
                    with st.expander(f"📄 {display_name}", expanded=False):
                        col_i, col_a = st.columns([3, 1])
                        with col_i:
                            st.write(f"**文件名:** {file_info['name']}")
                            st.write(f"**大小:** {format_file_size(file_info['size'])}")
                            st.write(f"**修改时间:** {format_timestamp(file_info['modified'])}")
                        with col_a:
                            delete_key = f"del_std_{file_info['name'].replace(' ', '_').replace('.', '_')}_{session_id}"
                            if st.button("🗑️ 删除", key=delete_key):
                                try:
                                    os.remove(file_info['path'])
                                    st.success(f"已删除: {file_info['name']}")
                                except Exception as e:
                                    st.error(f"删除失败: {e}")
            else:
                st.write("（未上传）")

        with tab_exam:
            exam_files = get_file_list(examined_dir)
            if exam_files:
                for file_info in exam_files:
                    display_name = truncate_filename(file_info['name'])
                    with st.expander(f"📄 {display_name}", expanded=False):
                        col_i, col_a = st.columns([3, 1])
                        with col_i:
                            st.write(f"**文件名:** {file_info['name']}")
                            st.write(f"**大小:** {format_file_size(file_info['size'])}")
                            st.write(f"**修改时间:** {format_timestamp(file_info['modified'])}")
                        with col_a:
                            delete_key = f"del_exam_{file_info['name'].replace(' ', '_').replace('.', '_')}_{session_id}"
                            if st.button("🗑️ 删除", key=delete_key):
                                try:
                                    os.remove(file_info['path'])
                                    st.success(f"已删除: {file_info['name']}")
                                except Exception as e:
                                    st.error(f"删除失败: {e}")
            else:
                st.write("（未上传）")

        with tab_results:
            # List files under generated/<session>/enterprise_standard_check/final_results
            final_dir = os.path.join(enterprise_out_root, 'final_results')
            if os.path.isdir(final_dir):
                final_files = get_file_list(final_dir)
                if final_files:
                    for file_info in final_files:
                        display_name = truncate_filename(file_info['name'])
                        with st.expander(f"📄 {display_name}", expanded=False):
                            col_i, col_a = st.columns([4, 1])
                            with col_i:
                                st.write(f"**文件名:** {file_info['name']}")
                                st.write(f"**大小:** {format_file_size(file_info['size'])}")
                                st.write(f"**修改时间:** {format_timestamp(file_info['modified'])}")
                            with col_a:
                                try:
                                    with open(file_info['path'], 'rb') as _fbin:
                                        _data = _fbin.read()
                                    st.download_button(
                                        label="⬇️ 下载",
                                        data=_data,
                                        file_name=file_info['name'],
                                        mime='application/octet-stream',
                                        key=f"dl_final_{file_info['name'].replace(' ', '_').replace('.', '_')}_{session_id}"
                                    )
                                except Exception as e:
                                    st.error(f"下载失败: {e}")
                                delete_key = f"del_final_{file_info['name'].replace(' ', '_').replace('.', '_')}_{session_id}"
                                if st.button("🗑️ 删除", key=delete_key):
                                    try:
                                        os.remove(file_info['path'])
                                        st.success(f"已删除: {file_info['name']}")
                                    except Exception as e:
                                        st.error(f"删除失败: {e}")
                else:
                    st.write("（暂无分析结果）")
            else:
                st.write("（暂无分析结果目录）")




    with col_main:
        # Two uploaders side by side
        col_std, col_exam = st.columns(2)
        with col_std:
            files_std = st.file_uploader("点击上传企业标准文件", type=None, accept_multiple_files=True, key=f"enterprise_std_{session_id}")
            if files_std:
                handle_file_upload(files_std, standards_dir)
                st.success(f"已上传 {len(files_std)} 个企业标准文件")
        with col_exam:
            files_exam = st.file_uploader("点击上传待检查文件", type=None, accept_multiple_files=True, key=f"enterprise_exam_{session_id}")
            if files_exam:
                handle_file_upload(files_exam, examined_dir)
                st.success(f"已上传 {len(files_exam)} 个待检查文件")

        # Start / Stop / Demo buttons
        btn_col1, btn_col_stop, btn_col2 = st.columns([1, 1, 1])
        with btn_col1:
            if st.button("开始", key=f"enterprise_start_button_{session_id}"):
                # Process PDFs (MinerU) and Word/PPT (Unstructured) into plain text
                area = st.container()
                with area:
                    # Step 0: Clean orphan .txt files that don't correspond to current uploads
                    try:
                        removed_std = _cleanup_orphan_txts(standards_dir, standards_txt_dir, st)
                        removed_exam = _cleanup_orphan_txts(examined_dir, examined_txt_dir, st)
                        if removed_std or removed_exam:
                            st.info(f"已清理无关文本 {removed_std + removed_exam} 个")
                    except Exception:
                        pass
                    # Step 0b: Clear previous run results so current run writes fresh outputs
                    try:
                        initial_dir = os.path.join(enterprise_out_root, 'initial_results')
                        os.makedirs(initial_dir, exist_ok=True)
                        cleared = 0
                        for fname in os.listdir(initial_dir):
                            fpath = os.path.join(initial_dir, fname)
                            if os.path.isfile(fpath):
                                try:
                                    os.remove(fpath)
                                    cleared += 1
                                except Exception:
                                    pass
                        if cleared:
                            st.info(f"已清空上次运行结果 {cleared} 个文件")
                    except Exception:
                        pass
                    st.markdown("**阅读企业标准文件中，10分钟左右，请等待...**")
                    created_std_pdf = _process_pdf_folder(standards_dir, standards_txt_dir, st, annotate_sources=True)
                    created_std_wp = _process_word_ppt_folder(standards_dir, standards_txt_dir, st, annotate_sources=True)
                    created_std_xls = _process_excel_folder(standards_dir, standards_txt_dir, st, annotate_sources=True)
                    # New: copy text-like files directly to standards_txt
                    created_std_text = _process_textlike_folder(standards_dir, standards_txt_dir, st)
                    # New: extract archives and process recursively
                    _ = _process_archives(standards_dir, standards_txt_dir, st)
                    st.markdown("**阅读待检查文件中，10分钟左右，请等待...**")
                    created_exam_pdf = _process_pdf_folder(examined_dir, examined_txt_dir, st, annotate_sources=False)
                    created_exam_wp = _process_word_ppt_folder(examined_dir, examined_txt_dir, st, annotate_sources=False)
                    created_exam_xls = _process_excel_folder(examined_dir, examined_txt_dir, st, annotate_sources=False)
                    # New: copy text-like files directly to examined_txt
                    created_exam_text = _process_textlike_folder(examined_dir, examined_txt_dir, st)
                    # New: extract archives and process recursively
                    _ = _process_archives(examined_dir, examined_txt_dir, st)

                    # If we have any txt, switch to running phase and rerun so streaming renders in main column
                    try:
                        std_txt_files = [f for f in os.listdir(standards_txt_dir) if f.lower().endswith('.txt')] if os.path.isdir(standards_txt_dir) else []
                        exam_txt_files = [f for f in os.listdir(examined_txt_dir) if f.lower().endswith('.txt')] if os.path.isdir(examined_txt_dir) else []
                        if not exam_txt_files:
                            st.warning("未发现待检查的 .txt 文本，跳过企业标准比对。")
                        else:
                            # --- Checkpoint preparation: generate prompts for all chunks and manifest ---
                            try:
                                checkpoint_dir = os.path.join(enterprise_out_root, 'checkpoint')
                                os.makedirs(checkpoint_dir, exist_ok=True)
                                # If previous manifest exists and all entries are done, clear checkpoint files
                                try:
                                    import json as _json
                                    _manifest_path = os.path.join(checkpoint_dir, 'manifest.json')
                                    if os.path.isfile(_manifest_path):
                                        with open(_manifest_path, 'r', encoding='utf-8') as _mf:
                                            _prev_manifest = _json.load(_mf) or {}
                                        _entries_prev = _prev_manifest.get('entries') or []
                                        _all_done = bool(_entries_prev) and all((str(e.get('status')) == 'done') for e in _entries_prev)
                                        if _all_done:
                                            for _fn in os.listdir(checkpoint_dir):
                                                _fp = os.path.join(checkpoint_dir, _fn)
                                                try:
                                                    if os.path.isfile(_fp):
                                                        os.remove(_fp)
                                                except Exception:
                                                    pass
                                except Exception:
                                    pass
                                # Build run_id based on current examined_txt files (name|size|mtime)
                                try:
                                    import hashlib
                                    infos = []
                                    for _n in sorted(exam_txt_files, key=lambda x: x.lower()):
                                        _pp = os.path.join(examined_txt_dir, _n)
                                        try:
                                            _st = os.stat(_pp)
                                            infos.append(f"{_n}|{_st.st_size}|{int(_st.st_mtime)}")
                                        except Exception:
                                            infos.append(f"{_n}|0|0")
                                    run_id = hashlib.sha1("\n".join(infos).encode('utf-8', errors='ignore')).hexdigest()
                                except Exception:
                                    run_id = ""
                                manifest = {"run_id": run_id, "entries": []}
                                # The same prompt prefix used later in streaming phase
                                prompt_prefix = (
                                    "请作为企业标准符合性检查专家，审阅待检查文件与企业标准是否一致。"
                                    "以列表形式列出不一致的点，并引用原文证据（简短摘录）、标明出处（提供企业标准文件的文件名）。\n"
                                    "输出的内容要言简意赅，列出不一致的点即可，不需要列出一致的点，也不需要列出企业标准中缺失的点，最后不需要总结。\n"
                                    "由于待检查文件较长，我将分成多个部分将其上传给你。以下是待检查文件的一部分。\n"
                                )
                                entry_id = 0
                                for _fname in sorted(exam_txt_files, key=lambda x: x.lower()):
                                    _src = os.path.join(examined_txt_dir, _fname)
                                    try:
                                        with open(_src, 'r', encoding='utf-8') as _f:
                                            _doc_text = _f.read()
                                    except Exception:
                                        _doc_text = ""
                                    _chunks = split_to_chunks(_doc_text, int(BISHENG_MAX_WORDS))
                                    for _ci, _piece in enumerate(_chunks):
                                        entry_id += 1
                                        _prompt_text = f"{prompt_prefix}{_piece}"
                                        # Use 1-based numbering per chunk within file for filenames, new convention
                                        _num = _ci + 1
                                        _prompt_name = f"checkpoint_prompt_{_fname}_pt{_num}.txt"
                                        _resp_name = f"checkpoint_response_{_fname}_pt{_num}.txt"
                                        _prompt_path = os.path.join(checkpoint_dir, _prompt_name)
                                        _response_path = os.path.join(checkpoint_dir, _resp_name)
                                        try:
                                            with open(_prompt_path, 'w', encoding='utf-8') as _pf:
                                                _pf.write(_prompt_text)
                                        except Exception:
                                            pass
                                        manifest["entries"].append({
                                            "id": entry_id,
                                            "file_name": _fname,
                                            "chunk_index": _ci,
                                            "prompt_file": _prompt_name,
                                            "response_file": _resp_name,
                                            "status": "not_done"
                                        })
                                # Write manifest.json atomically (paths stored as relative filenames for portability)
                                try:
                                    import json as _json, tempfile as _tmp, shutil as _sh
                                    _manifest_path = os.path.join(checkpoint_dir, 'manifest.json')
                                    with _tmp.NamedTemporaryFile('w', delete=False, encoding='utf-8', dir=checkpoint_dir) as _tf:
                                        _tf.write(_json.dumps(manifest, ensure_ascii=False, indent=2))
                                        _tmpname = _tf.name
                                    _sh.move(_tmpname, _manifest_path)
                                except Exception:
                                    pass
                            except Exception:
                                pass

                            st.session_state[f"enterprise_continue_running_{session_id}"] = False
                            st.session_state[f"enterprise_running_{session_id}"] = True
                            st.session_state[f"enterprise_std_txt_files_{session_id}"] = std_txt_files
                            st.session_state[f"enterprise_exam_txt_files_{session_id}"] = exam_txt_files
                            st.session_state[f"enterprise_out_root_{session_id}"] = enterprise_out_root
                            st.session_state[f"enterprise_standards_txt_dir_{session_id}"] = standards_txt_dir
                            st.session_state[f"enterprise_examined_txt_dir_{session_id}"] = examined_txt_dir
                            st.rerun()
                    except Exception as e:
                        st.error(f"企业标准比对流程异常：{e}")
                    
        with btn_col_stop:
            if st.button("停止", key=f"enterprise_stop_button_{session_id}"):
                st.session_state[f"enterprise_running_{session_id}"] = False
                try:
                    # Load current bisheng session id if any
                    bs_key = f"bisheng_session_{session_id}"
                    bisheng_sid = st.session_state.get(bs_key)
                    bs_cfg = CONFIG.get('bisheng', {})
                    base_url = (st.session_state.get(f"bisheng_{session_id}_base_url")
                                or os.getenv('BISHENG_BASE_URL') or bs_cfg.get('base_url') or 'http://10.31.60.11:3001')
                    stop_path = (st.session_state.get(f"bisheng_{session_id}_stop_path")
                                or os.getenv('BISHENG_STOP_PATH') or bs_cfg.get('stop_path') or '/api/v2/workflow/stop')
                    api_key = (st.session_state.get(f"bisheng_{session_id}_api_key")
                                or os.getenv('BISHENG_API_KEY') or bs_cfg.get('api_key') or '')
                    if not bisheng_sid:
                        st.info("当前无活动会话可停止。")
                    else:
                        res = stop_workflow(base_url, stop_path, bisheng_sid, api_key or None)
                        st.success(f"已请求停止，响应：{res}")
                except Exception as e:
                    st.error(f"停止失败：{e}")

            # Continue 按钮：始终可见。若无可续条目，提示信息；否则切换到运行分支消费 checkpoint。
            if st.button("继续", key=f"enterprise_continue_button_{session_id}"):
                if st.session_state.get(f"enterprise_running_{session_id}"):
                    st.info("当前流程正在运行，请先停止再继续。")
                else:
                    checkpoint_dir = os.path.join(enterprise_out_root, 'checkpoint')
                    manifest_path = os.path.join(checkpoint_dir, 'manifest.json')
                    _has_entries = False
                    try:
                        import json as _json
                        if os.path.isfile(manifest_path):
                            with open(manifest_path, 'r', encoding='utf-8') as _mf:
                                _m = _json.load(_mf) or {}
                                _ents = _m.get('entries') or []
                                _has_entries = len(_ents) > 0
                    except Exception:
                        _has_entries = False
                    if not _has_entries:
                        st.info("未发现跑到一半的项目")
                    else:
                        st.session_state[f"enterprise_running_{session_id}"] = False
                        st.session_state[f"enterprise_continue_running_{session_id}"] = True
                        st.rerun()

        with btn_col2:
            if st.button("演示", key=f"enterprise_demo_button_{session_id}"):
                # Copy demonstration files into the user's enterprise folders (no processing here)
                try:
                    import shutil
                    # Locate demonstration root (same convention as other tabs)
                    demo_base_dir = CONFIG["directories"]["cp_files"].parent / "demonstration"
                    demo_enterprise = os.path.join(str(demo_base_dir), "enterprise_standard_files")
                    # Subfolders to copy from → to
                    pairs = [
                        (os.path.join(demo_enterprise, "standards"), standards_dir),
                        (os.path.join(demo_enterprise, "examined_files"), examined_dir),
                        # New: copy demonstration prompt/response chunks into session enterprise output
                        # Entire folders copied under enterprise_out_root
                        (os.path.join(demo_enterprise, "prompt_text_chunks"), os.path.join(enterprise_out_root, "prompt_text_chunks")),
                        (os.path.join(demo_enterprise, "llm responses"), os.path.join(enterprise_out_root, "llm responses")),
                        # New: copy final_results for demo summary
                        (os.path.join(demo_enterprise, "final_results"), os.path.join(enterprise_out_root, "final_results")),
                        # New: copy pre-made prompted responses and json outputs for demo
                        (os.path.join(demo_enterprise, "prompted_llm responses_and_json"), os.path.join(enterprise_out_root, "prompted_llm responses_and_json")),
                    ]
                    files_copied = 0
                    for src, dst in pairs:
                        if not os.path.exists(src):
                            continue
                    # If source is a directory that we want to mirror (prompt_text_chunks / llm responses / final_results / prompted_llm responses_and_json)
                        if os.path.isdir(src) and (src.endswith("prompt_text_chunks") or src.endswith("llm responses") or src.endswith("final_results") or src.endswith("prompted_llm responses_and_json")):
                            os.makedirs(os.path.dirname(dst), exist_ok=True)
                            # Copy whole directory tree into enterprise_out_root subfolder
                            shutil.copytree(src, dst, dirs_exist_ok=True)
                            for root, _, files in os.walk(src):
                                files_copied += len([f for f in files if os.path.isfile(os.path.join(root, f))])
                            continue
                        # Otherwise treat as file list copy (standards / examined_files)
                        for name in os.listdir(src):
                            src_path = os.path.join(src, name)
                            dst_path = os.path.join(dst, name)
                            if os.path.isfile(src_path):
                                os.makedirs(dst, exist_ok=True)
                                shutil.copy2(src_path, dst_path)
                                files_copied += 1
                    # Trigger demo streaming phase
                    st.session_state[f"enterprise_demo_{session_id}"] = True
                    st.success(f"已复制演示文件：{files_copied} 个，开始演示…")
                except Exception as e:
                    st.error(f"演示文件复制失败: {e}")
                # Immediately rerun to render the demo streaming phase in main column
                st.rerun()

        # Render streaming phase in main column after rerun (mirrors special_symbols pattern)
        if st.session_state.get(f"enterprise_running_{session_id}"):
            # Retrieve context saved before rerun
            std_txt_files = st.session_state.get(f"enterprise_std_txt_files_{session_id}") or []
            exam_txt_files = st.session_state.get(f"enterprise_exam_txt_files_{session_id}") or []
            enterprise_out = st.session_state.get(f"enterprise_out_root_{session_id}") or enterprise_out_root
            std_txt_dir = st.session_state.get(f"enterprise_standards_txt_dir_{session_id}") or standards_txt_dir
            exam_txt_dir = st.session_state.get(f"enterprise_examined_txt_dir_{session_id}") or examined_txt_dir

            # Upload standards once (optional)
            if std_txt_files:
                with st.status("Sync standards to KB...", expanded=False) as status:
                    try:
                        kb_name_dyn = f"{session_id}_{TAB_SLUG}"
                        kid = find_knowledge_id_by_name(BISHENG_BASE_URL, BISHENG_API_KEY or None, kb_name_dyn)
                        if not kid:
                            kid = create_knowledge(BISHENG_BASE_URL, BISHENG_API_KEY or None, kb_name_dyn, model=str(KB_MODEL_ID))
                        if kid:
                            res = kb_sync_folder(
                                base_url=BISHENG_BASE_URL,
                                api_key=BISHENG_API_KEY or None,
                                knowledge_id=int(kid),
                                folder_path=std_txt_dir,
                                clear_first=False,
                                chunk_size=1000,
                                chunk_overlap=0,
                                separators=["\n\n", "\n"],
                                separator_rule=["after", "after"],
                            )
                            status.update(label=f"KB sync: uploaded {len(res.get('uploaded', []))}, deleted {len(res.get('deleted', []))}, skipped {len(res.get('skipped', []))}", state="complete")
                        else:
                            status.update(label="KB create/lookup failed (check server auth)", state="error")
                    except Exception as e:
                        status.update(label=f"KB sync failed: {e}", state="error")

            # Iterate examined texts
            exam_txt_files.sort(key=lambda x: x.lower())
            bisheng_session_id = st.session_state.get(f"bisheng_session_{session_id}")
            initial_dir = os.path.join(enterprise_out, 'initial_results')
            os.makedirs(initial_dir, exist_ok=True)
            # 预热步骤：在并发开始前，串行对 Bisheng Flow 发起一次极短请求，促使检索/LLM 初始化与缓存
            # 说明：首次请求常见的冷启动（模型加载、连接池、检索索引唤醒）会导致首批并发请求失败率上升；
            # 通过一次轻量的预热，可以显著降低“第一批全挂”的概率。返回内容无需使用。
            try:
                if not st.session_state.get(f"enterprise_warmup_done_{session_id}"):
                    warmup_prompt = "预热：请简短回复 'gotcha' 即可。"
                    _ = call_flow_process(
                        base_url=BISHENG_BASE_URL,
                        flow_id=BISHENG_FLOW_ID,
                        question=warmup_prompt,
                        kb_id=None,
                        input_node_id=FLOW_INPUT_NODE_ID,
                        api_key=BISHENG_API_KEY or None,
                        session_id=None,
                        history_count=0,
                        extra_tweaks={"CombineDocsChain-520ca": {"token_max": 5000}},
                        milvus_node_id=FLOW_MILVUS_NODE_ID,
                        es_node_id=FLOW_ES_NODE_ID,
                        timeout_s=60,
                        max_retries=0,
                        # clear_cache=True,
                    )
                    st.session_state[f"enterprise_warmup_done_{session_id}"] = True
            except Exception:
                pass
            # Start fresh: original per-file splitting and live prompting
            for idx_file, name in enumerate(exam_txt_files, start=1):
                src_path = os.path.join(exam_txt_dir, name)
                st.markdown(f"**📄 正在比对第{idx_file}个文件，共{len(exam_txt_files)}个：{name}**")
                try:
                    with open(src_path, 'r', encoding='utf-8') as f:
                        doc_text = f.read()
                except Exception as e:
                    st.error(f"读取失败：{e}")
                    continue
                if not doc_text.strip():
                    st.info("文件为空，跳过。")
                    continue
                chunks = split_to_chunks(doc_text, int(BISHENG_MAX_WORDS))
                prompt_prefix = (
                    "请作为企业标准符合性检查专家，审阅待检查文件与企业标准是否一致。"
                    "以列表形式列出不一致的点，并引用原文证据（简短摘录）、标明出处（提供企业标准文件的文件名）。\n"
                    "输出的内容要言简意赅，列出不一致的点即可，不需要列出一致的点，也不需要列出企业标准中缺失的点，最后不需要总结。\n"
                    "由于待检查文件较长，我将分成多个部分将其上传给你。以下是待检查文件的一部分。\n"
                )
                full_out_text = ""
                prompt_texts = []
                for i, piece in enumerate(chunks, start=1):
                    col_prompt, col_response = st.columns([1, 1])
                    prompt_text = f"{prompt_prefix}{piece}"
                    prompt_texts.append(prompt_text)
                    with col_prompt:
                        st.markdown(f"提示词（第{i}部分，共{len(chunks)}部分）")
                        prompt_container = st.container(height=400)
                        with prompt_container:
                            with st.chat_message("user"):
                                prompt_placeholder = st.empty()
                                _stream_text(prompt_placeholder, prompt_text, render_method="text")
                            st.chat_input(placeholder="", disabled=True, key=f"enterprise_prompt_{session_id}_{idx_file}_{i}")
                    with col_response:
                        st.markdown(f"AI比对结果（第{i}部分，共{len(chunks)}部分）")
                        response_container = st.container(height=400)
                        with response_container:
                            with st.chat_message("assistant"):
                                response_placeholder = st.empty()
                                try:
                                    # Determine per-user KB name and call Flow with tweaks for per-run KB binding
                                    kb_name_dyn = f"{session_id}_{TAB_SLUG}"
                                    kid = find_knowledge_id_by_name(BISHENG_BASE_URL, BISHENG_API_KEY or None, kb_name_dyn)
                                    start_ts = time.time()
                                    res = call_flow_process(
                                        base_url=BISHENG_BASE_URL,
                                        flow_id=BISHENG_FLOW_ID,
                                        question=prompt_text,
                                        kb_id=kid,
                                        input_node_id=FLOW_INPUT_NODE_ID,
                                        api_key=BISHENG_API_KEY or None,
                                        session_id=bisheng_session_id,
                                        history_count=0,
                                        extra_tweaks=None,
                                        milvus_node_id=FLOW_MILVUS_NODE_ID,
                                        es_node_id=FLOW_ES_NODE_ID,
                                        timeout_s=180,        # 新增
                                        max_retries=2,        # 新增
                                        # clear_cache=True,
                                        )
                                    ans_text, new_sid = parse_flow_answer(res)
                                    dur_ms = int((time.time() - start_ts) * 1000)
                                    try:
                                        from datetime import datetime as _dt
                                        _log_llm_metrics(
                                            enterprise_out,
                                            session_id,
                                            {
                                                "ts": _dt.now().isoformat(timespec="seconds"),
                                                "engine": "bisheng",
                                                "model": "qwen3",
                                                "session_id": bisheng_session_id or "",
                                                "file": name,
                                                "part": i,
                                                "phase": "compare",
                                                "prompt_chars": len(prompt_text or ""),
                                                "prompt_tokens": _estimate_tokens(prompt_text or ""),
                                                "output_chars": len(ans_text or ""),
                                                "output_tokens": _estimate_tokens(ans_text or ""),
                                                "duration_ms": dur_ms,
                                                "success": 1 if (ans_text or "").strip() else 0,
                                                "error": res.get("error") if isinstance(res, dict) else "",
                                            },
                                        )
                                    except Exception:
                                        pass
                                    if new_sid:
                                        bisheng_session_id = new_sid
                                        st.session_state[f"bisheng_session_{session_id}"] = bisheng_session_id
                                    response_placeholder.write(ans_text or "")
                                    full_out_text += ("\n\n" if full_out_text else "") + (ans_text or "")
                                    # Mark progress in checkpoint only if Bisheng output contains <think>
                                    if '<think>' in (ans_text or ''):
                                        try:
                                            checkpoint_dir = os.path.join(enterprise_out, 'checkpoint')
                                            os.makedirs(checkpoint_dir, exist_ok=True)
                                            resp_fname = f"checkpoint_response_{name}_pt{i}.txt"
                                            resp_path = os.path.join(checkpoint_dir, resp_fname)
                                            # atomic write: temp file + move
                                            import tempfile as _tmp, shutil as _sh
                                            with _tmp.NamedTemporaryFile('w', delete=False, encoding='utf-8', dir=checkpoint_dir) as _tf:
                                                _tf.write(ans_text or "")
                                                _tmpname = _tf.name
                                            _sh.move(_tmpname, resp_path)
                                            # load and update manifest
                                            import json as _json, tempfile as _tmp, shutil as _sh
                                            m_path = os.path.join(checkpoint_dir, 'manifest.json')
                                            _m = None
                                            try:
                                                with open(m_path, 'r', encoding='utf-8') as _mf:
                                                    _m = _json.load(_mf) or {}
                                            except Exception:
                                                _m = None
                                            if isinstance(_m, dict) and isinstance(_m.get('entries'), list):
                                                for __e in _m['entries']:
                                                    if __e.get('file_name') == name and int(__e.get('chunk_index', -1)) == (i - 1):
                                                        __e['status'] = 'done'
                                                        break
                                                with _tmp.NamedTemporaryFile('w', delete=False, encoding='utf-8', dir=checkpoint_dir) as _tf:
                                                    _tf.write(_json.dumps(_m, ensure_ascii=False, indent=2))
                                                    _tmpname = _tf.name
                                                _sh.move(_tmpname, m_path)
                                        except Exception:
                                            pass
                                except Exception as e:
                                    response_placeholder.error(f"调用失败：{e}")
                            st.chat_input(placeholder="", disabled=True, key=f"enterprise_response_{session_id}_{idx_file}_{i}")
                # Persist per-file combined output
                try:
                    name_no_ext = os.path.splitext(name)[0]
                    _persist_compare_outputs(initial_dir, name_no_ext, prompt_texts, full_out_text)
                    _summarize_with_ollama(initial_dir, enterprise_out, session_id, name_no_ext, full_out_text)
                except Exception as e:
                    st.error(f"保存结果失败：{e}")


            # End of current run; clear running flag (no final aggregation here)
            try:
                st.session_state[f"enterprise_running_{session_id}"] = False
            except Exception as e:
                st.error(f"流程收尾失败：{e}")

            # After LLM step: aggregate json_*_ptN.txt into CSV and XLSX in final_results
            try:
                _aggregate_outputs(initial_dir, enterprise_out, session_id)
            except Exception as e:
                st.error(f"汇总导出失败：{e}")

        # Demo streaming phase (reads from prepared prompt/response chunks; no LLM calls)
        if st.session_state.get(f"enterprise_demo_{session_id}"):
            # Directories prepared by demo button copy
            prompt_dir = os.path.join(enterprise_out_root, 'prompt_text_chunks')
            resp_dir = os.path.join(enterprise_out_root, 'llm responses')
            final_dir = os.path.join(enterprise_out_root, 'final_results')
            prompted_and_json_dir = os.path.join(enterprise_out_root, 'prompted_llm responses_and_json')
            # Collect prompt chunk files
            prompt_files = []
            try:
                if os.path.isdir(prompt_dir):
                    for f in os.listdir(prompt_dir):
                        if f.lower().endswith('.txt'):
                            prompt_files.append(f)
            except Exception:
                prompt_files = []
            # Natural sort by base name and numeric part index
            _prompt_entries = []
            for _f in prompt_files:
                _m = re.match(r"^(?P<base>.+)_pt(?P<idx>\d+)\.txt$", _f)
                if _m:
                    _prompt_entries.append((_m.group('base').lower(), int(_m.group('idx')), _f))
                else:
                    _prompt_entries.append(("", 0, _f))
            _prompt_entries.sort(key=lambda t: (t[0], t[1]))
            # Render each prompt/response pair in UI (original demo)
            for _, _, fname in _prompt_entries:
                m = re.match(r"^(?P<base>.+)_pt(?P<idx>\d+)\.txt$", fname)
                if not m:
                    continue
                base = m.group('base')
                idx = m.group('idx')
                prompt_path = os.path.join(prompt_dir, fname)
                resp_name = f"response_{base}_pt{idx}.txt"
                resp_path = os.path.join(resp_dir, resp_name)
                # Read prompt content
                try:
                    with open(prompt_path, 'r', encoding='utf-8') as f:
                        prompt_text = f.read()
                except Exception:
                    prompt_text = ""
                # Read response content (optional)
                resp_text = None
                if os.path.isfile(resp_path):
                    try:
                        with open(resp_path, 'r', encoding='utf-8') as f:
                            resp_text = f.read()
                    except Exception:
                        resp_text = None
                col_prompt, col_response = st.columns([1, 1])
                with col_prompt:
                    st.markdown(f"提示词（{base} - 第{idx}部分）")
                    prompt_container = st.container(height=400)
                    with prompt_container:
                        with st.chat_message("user"):
                            prompt_placeholder = st.empty()
                            _stream_text(prompt_placeholder, prompt_text, render_method="text", delay=0.1)
                        st.chat_input(placeholder="", disabled=True, key=f"enterprise_demo_prompt_{session_id}_{base}_{idx}")
                with col_response:
                    st.markdown(f"示例比对结果（{base} - 第{idx}部分）")
                    response_container = st.container(height=400)
                    with response_container:
                        with st.chat_message("assistant"):
                            resp_placeholder = st.empty()
                            if resp_text is None:
                                resp_placeholder.info("未找到对应示例结果。")
                            else:
                                _stream_text(resp_placeholder, resp_text, render_method="write", delay=0.1)
                            st.chat_input(placeholder="", disabled=True, key=f"enterprise_demo_resp_{session_id}_{base}_{idx}")
            # (Removed hardcoded final report section for demo)
            # End of demo streaming pass; reset the flag
                st.session_state[f"enterprise_demo_{session_id}"] = False

            # New demo rendering: prompted_response_* and corresponding json_* from pre-made folder
            try:
                if os.path.isdir(prompted_and_json_dir):
                    # Find prompted_response_*.txt files and for each, display prompt and its json_ response
                    demo_parts = [f for f in os.listdir(prompted_and_json_dir) if f.startswith('prompted_response_') and f.lower().endswith('.txt')]
                    # Natural sort by extracted name and numeric part index
                    _entries = []
                    for _pf in demo_parts:
                        _m = re.match(r"^prompted_response_(?P<name>.+)_pt(?P<idx>\d+)\.txt$", _pf)
                        if _m:
                            _entries.append((_m.group('name').lower(), int(_m.group('idx')), _pf))
                        else:
                            _entries.append(("", 0, _pf))
                    _entries.sort(key=lambda t: (t[0], t[1]))
                    total_parts = len(_entries)
                    for idx, (_, _, pf) in enumerate(_entries, start=1):
                        p_path = os.path.join(prompted_and_json_dir, pf)
                        # Map to json_<name>.txt in same folder
                        json_name = f"json_{pf[len('prompted_response_'):] }"
                        j_path = os.path.join(prompted_and_json_dir, json_name)
                        # Read prompt
                        try:
                            with open(p_path, 'r', encoding='utf-8') as f:
                                ptext = f.read()
                        except Exception:
                            ptext = ""
                        # Read json response (text form)
                        try:
                            with open(j_path, 'r', encoding='utf-8') as f:
                                jtext = f.read()
                        except Exception:
                            jtext = ""
                        # Two column display
                        col_lp, col_lr = st.columns([1, 1])
                        with col_lp:
                            st.markdown(f"生成汇总表格提示词（第{idx}部分，共{total_parts}部分）")
                            pc = st.container(height=400)
                            with pc:
                                with st.chat_message("user"):
                                    ph = st.empty()
                                    _stream_text(ph, ptext, render_method="text", delay=0.1)
                            st.chat_input(placeholder="", disabled=True, key=f"enterprise_demo_prompted_prompt_{session_id}_{idx}")
                        with col_lr:
                            st.markdown(f"生成汇总表格结果（第{idx}部分，共{total_parts}部分）")
                            rc = st.container(height=400)
                            with rc:
                                with st.chat_message("assistant"):
                                    ph2 = st.empty()
                                    _stream_text(ph2, jtext, render_method="write", delay=0.1)
                            st.chat_input(placeholder="", disabled=True, key=f"enterprise_demo_prompted_resp_{session_id}_{idx}")
            except Exception:
                pass

            # Add demo download buttons for CSV/XLSX in final_results
            try:
                if os.path.isdir(final_dir):
                    csv_files = [f for f in os.listdir(final_dir) if f.lower().endswith('.csv')]
                    xlsx_files = [f for f in os.listdir(final_dir) if f.lower().endswith('.xlsx')]
                    def _latest(path_list):
                        if not path_list:
                            return None
                        paths = [os.path.join(final_dir, f) for f in path_list]
                        paths.sort(key=lambda p: os.path.getmtime(p))
                        return paths[-1]
                    latest_csv = _latest(csv_files)
                    latest_xlsx = _latest(xlsx_files)
                    if latest_csv:
                        with open(latest_csv, 'rb') as fcsv:
                            st.download_button(label="下载CSV结果(演示)", data=fcsv.read(), file_name=os.path.basename(latest_csv), mime='text/csv', key=f"demo_download_csv_{session_id}")
                    if latest_xlsx:
                        with open(latest_xlsx, 'rb') as fxlsx:
                            st.download_button(label="下载Excel结果(演示)", data=fxlsx.read(), file_name=os.path.basename(latest_xlsx), mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', key=f"demo_download_xlsx_{session_id}")
            except Exception:
                pass

        # Continue streaming/processing section (consume checkpoint: done → playback; not_done → run)
        if st.session_state.get(f"enterprise_continue_running_{session_id}"):
            # Retrieve context and init dirs
            std_txt_files = st.session_state.get(f"enterprise_std_txt_files_{session_id}") or []
            exam_txt_files = st.session_state.get(f"enterprise_exam_txt_files_{session_id}") or []
            enterprise_out = st.session_state.get(f"enterprise_out_root_{session_id}") or enterprise_out_root
            initial_dir = os.path.join(enterprise_out, 'initial_results')
            os.makedirs(initial_dir, exist_ok=True)
            checkpoint_dir = os.path.join(enterprise_out, 'checkpoint')
            manifest_path = os.path.join(checkpoint_dir, 'manifest.json')
            _manifest = None
            try:
                import json as _json
                with open(manifest_path, 'r', encoding='utf-8') as _mf:
                    _manifest = _json.load(_mf) or {}
            except Exception:
                _manifest = None
            if not _manifest or not isinstance(_manifest.get('entries'), list) or not _manifest['entries']:
                st.info("未发现跑到一半的项目")
                st.session_state[f"enterprise_continue_running_{session_id}"] = False
            else:
                # group by file in entry order
                from collections import OrderedDict as _OD
                _entries_sorted = sorted((_manifest.get('entries') or []), key=lambda e: int(e.get('id', 0)))
                _groups = _OD()
                for _e in _entries_sorted:
                    fname = str(_e.get('file_name', ''))
                    _groups.setdefault(fname, []).append(_e)
                bisheng_session_id = st.session_state.get(f"bisheng_session_{session_id}")
                for name, _elist in _groups.items():
                    st.markdown(f"**📄 继续比对：{name}**")
                    full_out_text = ""
                    prompt_texts = []
                    _total_parts = len(_elist)
                    try:
                        _elist.sort(key=lambda e: (int(e.get('chunk_index', 0)), int(e.get('id', 0))))
                    except Exception:
                        pass
                    for _i, _entry in enumerate(_elist, start=1):
                        _prompt_file = str(_entry.get('prompt_file', ''))
                        _response_file = str(_entry.get('response_file', ''))
                        _status = str(_entry.get('status', 'not_done'))
                        _prompt_path = os.path.join(checkpoint_dir, _prompt_file)
                        _response_path = os.path.join(checkpoint_dir, _response_file)
                        # load prompt
                        try:
                            with open(_prompt_path, 'r', encoding='utf-8') as _pf:
                                prompt_text = _pf.read()
                        except Exception:
                            prompt_text = ""
                        prompt_texts.append(prompt_text)
                        col_prompt, col_response = st.columns([1, 1])
                        with col_prompt:
                            st.markdown(f"提示词（第{_i}部分，共{_total_parts}部分）")
                            prompt_container = st.container(height=400)
                            with prompt_container:
                                with st.chat_message("user"):
                                    prompt_placeholder = st.empty()
                                    _stream_text(prompt_placeholder, prompt_text, render_method="text")
                            st.chat_input(placeholder="", disabled=True, key=f"enterprise_continue_prompt_{session_id}_{name}_{_i}")
                        with col_response:
                            st.markdown(f"AI比对结果（第{_i}部分，共{_total_parts}部分）")
                            response_container = st.container(height=400)
                            with response_container:
                                with st.chat_message("assistant"):
                                    ph = st.empty()
                                    if _status == 'done':
                                        # playback response
                                        try:
                                            with open(_response_path, 'r', encoding='utf-8') as _rf:
                                                _resp_text = _rf.read()
                                        except Exception:
                                            _resp_text = ""
                                        _stream_text(ph, _resp_text, render_method="write", delay=0.1)
                                        full_out_text += ("\n\n" if full_out_text else "") + (_resp_text or "")
                                    else:
                                        # run LLM and update manifest
                                        try:
                                            kb_name_dyn = f"{session_id}_{TAB_SLUG}"
                                            kid = find_knowledge_id_by_name(BISHENG_BASE_URL, BISHENG_API_KEY or None, kb_name_dyn)
                                            response_text = ""
                                            res = call_flow_process(
                                                base_url=BISHENG_BASE_URL,
                                                flow_id=BISHENG_FLOW_ID,
                                                question=prompt_text,
                                                kb_id=kid,
                                                input_node_id=FLOW_INPUT_NODE_ID,
                                                api_key=BISHENG_API_KEY or None,
                                                session_id=bisheng_session_id,
                                                history_count=0,
                                                extra_tweaks=None,
                                                milvus_node_id=FLOW_MILVUS_NODE_ID,
                                                es_node_id=FLOW_ES_NODE_ID,
                                                timeout_s=180,
                                                max_retries=2,
                                            )
                                            ans_text, new_sid = parse_flow_answer(res)
                                            if new_sid:
                                                bisheng_session_id = new_sid
                                                st.session_state[f"bisheng_session_{session_id}"] = bisheng_session_id
                                            ph.write(ans_text or "")
                                            full_out_text += ("\n\n" if full_out_text else "") + (ans_text or "")
                                            # save resp and mark done (atomic writes like Start) only if Bisheng output contains <think>
                                            if '<think>' in (ans_text or ''):
                                                try:
                                                    import tempfile as _tmp, shutil as _sh
                                                    with _tmp.NamedTemporaryFile('w', delete=False, encoding='utf-8', dir=checkpoint_dir) as _tf:
                                                        _tf.write(ans_text or "")
                                                        _tmpname = _tf.name
                                                    _sh.move(_tmpname, _response_path)
                                                except Exception:
                                                    pass
                                                try:
                                                    import json as _json, tempfile as _tmp, shutil as _sh
                                                    for __e in _manifest.get('entries', []):
                                                        if (
                                                            str(__e.get('file_name')) == str(name)
                                                            and int(__e.get('chunk_index', -1)) == int(_entry.get('chunk_index', -1))
                                                        ):
                                                            __e['status'] = 'done'
                                                            break
                                                    with _tmp.NamedTemporaryFile('w', delete=False, encoding='utf-8', dir=checkpoint_dir) as _tf:
                                                        _tf.write(_json.dumps(_manifest, ensure_ascii=False, indent=2))
                                                        _tmpname = _tf.name
                                                    _sh.move(_tmpname, manifest_path)
                                                except Exception:
                                                    pass
                                        except Exception as e:
                                            ph.error(f"调用失败：{e}")
                            st.chat_input(placeholder="", disabled=True, key=f"enterprise_continue_response_{session_id}_{name}_{_i}")
                # write combined prompt/response for this file
                    name_no_ext = os.path.splitext(name)[0]
                    try:
                        _persist_compare_outputs(initial_dir, name_no_ext, prompt_texts, full_out_text)
                    except Exception:
                        pass
                    # summarize for this file as well
                    try:
                        _summarize_with_ollama(initial_dir, enterprise_out, session_id, name_no_ext, full_out_text)
                    except Exception:
                        pass
            # End continue branch
            # After continue: aggregate all outputs to CSV/XLSX/Word like Start does
            try:
                _aggregate_outputs(initial_dir, enterprise_out, session_id)
            except Exception as e:
                st.error(f"汇总导出失败：{e}")
# The end
